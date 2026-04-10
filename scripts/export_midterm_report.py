#!/usr/bin/env python3
"""한양대 부동산융합대학원 석사학위 논문작성 중간보고서 DOCX 생성
PDF 샘플의 레이아웃을 재현하여 현재 논문(서울시 아파트 매매가격 XAI 분석)으로 작성.
"""
import os
from docx import Document
from docx.shared import Pt, Cm, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

BASE_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), '..')
OUT = os.path.join(BASE_DIR, 'paper', '석사학위논문작성중간보고서_박현근.docx')

FONT = '맑은 고딕'

def set_font(run, size=11, bold=False, font=FONT, color=None):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = font
    if color:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{font}"/>')
        rPr.append(rFonts)
    else:
        rFonts.set(qn('w:eastAsia'), font)

def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)

def set_cell_borders(cell, size='4'):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="{size}" w:space="0" w:color="000000"/>'
        f'<w:left w:val="single" w:sz="{size}" w:space="0" w:color="000000"/>'
        f'<w:bottom w:val="single" w:sz="{size}" w:space="0" w:color="000000"/>'
        f'<w:right w:val="single" w:sz="{size}" w:space="0" w:color="000000"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(tcBorders)

def add_heading_bar(doc, text):
    """■ 섹션 제목 (파란 막대 느낌)"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run('■ ' + text)
    set_font(r, size=14, bold=True, color=RGBColor(0x0B, 0x2E, 0x6E))
    # bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="12" w:space="1" w:color="0B2E6E"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)

def add_box_paragraph(doc, lines, heading=None):
    """테두리가 있는 본문 박스"""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    tbl.columns[0].width = Cm(16)
    cell = tbl.rows[0].cells[0]
    cell.width = Cm(16)
    set_cell_borders(cell, size='6')
    cell.paragraphs[0].text = ''
    if heading:
        p0 = cell.paragraphs[0]
        r = p0.add_run(heading)
        set_font(r, 12, bold=True)
        p0.paragraph_format.space_after = Pt(6)
    for line in lines:
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.4
        r = p.add_run(line)
        set_font(r, 10.5)
    # padding
    tcPr = cell._tc.get_or_add_tcPr()
    mar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="180" w:type="dxa"/>'
        f'<w:left w:w="200" w:type="dxa"/>'
        f'<w:bottom w:w="180" w:type="dxa"/>'
        f'<w:right w:w="200" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(mar)

def main():
    doc = Document()

    # --- margins ---
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # --- 제목 ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run('석사학위 논문작성 중간보고서')
    set_font(r, size=22, bold=True, color=RGBColor(0x0B, 0x2E, 0x6E))
    # underline bar
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="18" w:space="2" w:color="0B2E6E"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)

    # --- 제출자 섹션 ---
    add_heading_bar(doc, '제 출 자')

    tbl = doc.add_table(rows=2, cols=5)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    widths = [Cm(3.6), Cm(2.0), Cm(3.8), Cm(2.8), Cm(3.8)]
    headers = ['전공명', '기수', '학번', '성명', '연락처']
    values = ['빅데이터', '82', '2022151513', '박현근', '010-9592-0923']
    for i, w in enumerate(widths):
        tbl.columns[i].width = w
    for i, h in enumerate(headers):
        c = tbl.rows[0].cells[i]
        c.width = widths[i]
        set_cell_borders(c, '8')
        set_cell_bg(c, 'D9D9D9')
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_font(r, 11, bold=True)
    for i, v in enumerate(values):
        c = tbl.rows[1].cells[i]
        c.width = widths[i]
        set_cell_borders(c, '8')
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(v)
        set_font(r, 11)
    # row height
    for row in tbl.rows:
        row.height = Cm(0.9)

    # --- 보고 내용 섹션 ---
    add_heading_bar(doc, '보고 내용')

    # 박스: 1. 논문개요
    overview_lines = [
        '"XGBoost와 SHAP을 활용한 서울시 아파트 매매가격 결정요인 분석"',
        '',
        '서울시 25개 자치구 215개 행정동의 아파트 매매 실거래가 391,826건(2019.1~2025.12)을 대상으로 OLS·Random Forest·XGBoost 3개 모형의 예측 성능을 비교하고, SHAP(SHapley Additive exPlanations) 기법으로 주택 가격 결정요인의 영향력을 해석한다.',
        '전용면적·건물연령·층 등 주택 특성 변수와 강남구분·지하철역·학교·백화점·학원·어린이집·CCTV·공원·도서관 등 입지/생활 인프라 변수, 기준금리·CD금리·소비자물가지수·M2 등 거시경제 변수를 함께 투입하여 지역 이질성과 시장 국면 효과까지 포괄하는 통합 실증 설계를 수행한다.',
    ]
    add_box_paragraph(doc, overview_lines, heading='1. 논문개요')

    # 박스: 2. 논문작성 진행사항
    progress_lines = [
        '1. 데이터 수집·구축 완료',
        '  * 국토교통부 실거래가 API로 2019~2025년 서울 아파트 매매 391,826건 수집',
        '  * 서울열린데이터·NEIS·한국은행 ECOS API로 지하철·공원·학교·대규모점포·거시경제 지표 결합',
        '  * 건물연령·강남구분 등 파생변수 생성 및 행정동 단위(215개) 집계',
        '',
        '2. 예측 모형 적합 및 비교 완료',
        '  * OLS 다중회귀, Random Forest, XGBoost 3종을 동일 자료·동일 하이퍼파라미터 하에서 적합',
        '  * 무작위 분할 기준 XGBoost R² 0.970, RMSE 13,765만원, MAPE 7.65%로 OLS 대비 RMSE 72% 감소',
        '',
        '3. SHAP 기반 설명가능성 분석 완료',
        '  * 변수 중요도 및 의존성 플롯 산출, 전용면적·강남구분·건물연령·M2·백화점수가 상위 5대 요인으로 확인',
        '',
        '4. 강건성 검증(4종) 완료',
        '  * 시간순 분할(Train 2019~2023, Test 2024~2025)로 국면 이동에 대한 일반화 성능 점검',
        '  * 법정동·아파트명 기준 GroupKFold 분할(8,601 단지, overlap 0)로 단지 반복 누수 정량화 → XGB R² 0.846',
        '  * 학원수·어린이집수 Ablation으로 구 단위 프록시의 한계 기여 ΔR²≈−0.0003 확인',
        '  * 강남/비강남 지역별 별도 모형 적합으로 지역 이질성 재확인(학원수 비강남 0.70배, 어린이집수 0.12배, 백화점수 강남 5.29배)',
        '',
        '5. 논문 본문 초안 작성 완료',
        '  * 국문 5장 구조(서론·이론·방법·결과·결론) 및 국·영문 초록 작성, 표 15개·그림 9개 포함',
        '  * 지도교수 및 내부 재리뷰 반영, 향후 최종 심사 제출 준비 단계',
    ]
    add_box_paragraph(doc, progress_lines, heading='2. 논문작성 진행사항')

    # 박스: 3. 지도교수 의견 및 평가
    advisor_lines = [
        '',
        '',
        '',
        '',
    ]
    add_box_paragraph(doc, advisor_lines, heading='3. 지도교수 의견 및 평가 (반드시 기재 요망)')

    # --- 날짜 ---
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    r = p.add_run('2026 년  04 월  10일')
    set_font(r, 12, bold=True)

    # --- 서명 ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(16)
    r = p.add_run('제 출 자 :    박 현 근    (인)')
    set_font(r, 12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(6)
    r = p.add_run('지 도 교 수 :                   (인)')
    set_font(r, 12)

    # --- 수신 ---
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    r = p.add_run('부동산융합대학원장 귀하')
    set_font(r, 14, bold=True)

    doc.save(OUT)
    size = os.path.getsize(OUT) // 1024
    print(f'완료: {OUT}')
    print(f'크기: {size}KB')

if __name__ == '__main__':
    main()
