#!/usr/bin/env python3
"""
논문 DOCX 변환 v2 — 한양대학교 부동산융합대학원 석사논문 공식 서식
출처: https://gupd.hanyang.ac.kr/front/ko/bachelor/thesis
참고: 조민지(2023) 한양대 석사논문

편집용지: A4
워드 여백: 위 4cm, 아래 4cm, 좌 3.5cm, 우 3.5cm, 머리말 1.5cm, 꼬리말 1.5cm
본문: 10~11pt 바탕체, 줄간격 160~200%
큰제목(장): 16pt 진하게 / 중간제목(절): 13pt 진하게
"""
import os, re, io, struct
from datetime import datetime
import latex2mathml.converter as _l2m
import mathml2omml as _m2o
from docx import Document

OMML_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'

from docx.shared import Pt, Cm, Mm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

BASE_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), '..')
PAPER_DIR = os.path.join(BASE_DIR, 'paper')

# 한양대 작성 지침은 명조·신명조·바탕체 등 일반 폰트를 권장한다.
# 참고 표지의 명조계 조판과 Windows Word 변환 안정성을 함께 맞추기 위해 HY신명조로 통일한다.
FONT_KR = 'HY신명조'


# 본문 폭은 A4 좌우 여백 3.5cm 기준 약 14cm이다.
# 그림은 캡션과 함께 한 문단으로 묶이므로, 큰 그림이 하단에서 밀리며
# 과도한 공백을 만들지 않도록 그림별 최대 폭/높이를 제한한다.
IMAGE_SIZE_LIMITS_CM = {
    'fig1_research_flow.png': (12.0, 13.3),
    'fig2_xgboost_concept.png': (13.6, 8.8),
    'fig3_shap_framework.png': (13.6, 10.7),
    'fig4_shap_bar.png': (11.0, 8.8),
    'fig5_shap_summary.png': (9.8, 9.9),
    'fig6_dep_건물연령.png': (10.4, 7.4),
    'fig7_dep_childcare_count_1000m.png': (10.4, 7.4),
    'fig8_dep_subway_nearest.png': (10.4, 7.4),
    'fig9_dep_department_nearest.png': (10.4, 7.4),
    'fig10_ablation.png': (13.6, 8.2),
    'fig11_region_shap.png': (13.6, 8.2),
    'fig12_year_region_heatmap.png': (11.0, 7.3),
    'fig13_top1_timeline.png': (13.8, 6.3),
}
DEFAULT_IMAGE_LIMIT_CM = (12.0, 9.0)


def png_dimensions(path):
    with open(path, 'rb') as f:
        header = f.read(24)
    if header[:8] != b'\x89PNG\r\n\x1a\n':
        return None
    return struct.unpack('>II', header[16:24])


def image_width_cm_for_docx(path):
    max_w, max_h = IMAGE_SIZE_LIMITS_CM.get(
        os.path.basename(path),
        DEFAULT_IMAGE_LIMIT_CM,
    )
    dims = png_dimensions(path)
    if not dims:
        return max_w
    px_w, px_h = dims
    if px_w <= 0:
        return max_w
    ratio = px_h / px_w
    width = max_w
    if width * ratio > max_h:
        width = max_h / ratio
    return max(7.5, min(width, max_w))


def set_font(run, size=11, bold=False, font=FONT_KR, italic=False):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.name = font
    run.font.color.rgb = RGBColor(0, 0, 0)  # 학위논문 본문·헤딩 모두 검정
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(
            f'<w:rFonts {nsdecls("w")} '
            f'w:ascii="{font}" w:hAnsi="{font}" w:eastAsia="{font}" w:cs="{font}"/>'
        )
        rPr.append(rFonts)
    else:
        rFonts.set(qn('w:eastAsia'), font)
        rFonts.set(qn('w:ascii'), font)
        rFonts.set(qn('w:hAnsi'), font)
        rFonts.set(qn('w:cs'), font)


def set_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_width(cell, width_cm):
    tcPr = cell._tc.get_or_add_tcPr()
    for ex in list(tcPr.findall(qn('w:tcW'))):
        tcPr.remove(ex)
    tcPr.append(parse_xml(
        f'<w:tcW {nsdecls("w")} w:w="{int(Cm(width_cm).twips)}" w:type="dxa"/>'
    ))


def set_cell_margins(cell, top=70, left=90, bottom=70, right=90):
    tcPr = cell._tc.get_or_add_tcPr()
    for ex in list(tcPr.findall(qn('w:tcMar'))):
        tcPr.remove(ex)
    tcPr.append(parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    ))


def set_table_grid(tbl, widths_cm):
    tbl.autofit = False
    tblPr = tbl._tbl.tblPr
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblPr.append(parse_xml(
            f'<w:tblW {nsdecls("w")} w:w="{int(Cm(sum(widths_cm)).twips)}" w:type="dxa"/>'
        ))
    else:
        tblW.set(qn('w:w'), str(int(Cm(sum(widths_cm)).twips)))
        tblW.set(qn('w:type'), 'dxa')

    for ex in list(tbl._tbl.findall(qn('w:tblGrid'))):
        tbl._tbl.remove(ex)
    grid_cols = ''.join(
        f'<w:gridCol w:w="{int(Cm(w).twips)}"/>' for w in widths_cm
    )
    tbl._tbl.insert(1, parse_xml(f'<w:tblGrid {nsdecls("w")}>{grid_cols}</w:tblGrid>'))

    for row in tbl.rows:
        for idx, cell in enumerate(row.cells):
            if idx < len(widths_cm):
                set_cell_width(cell, widths_cm[idx])


def add_page_number(section, start_num=None, fmt='decimal', placeholder_fmt=None):
    """페이지 번호 추가 (하단 중앙, "- N -" 형식). fmt='decimal'|'lowerRoman'.

    placeholder_fmt: LibreOffice 호환을 위해 sectPr에 pgNumType을 2개 두는데,
    placeholder는 reset되지 않은 fmt이고 그 다음 element가 reset.
    fmt와 placeholder_fmt가 다를 때만 LibreOffice가 reset 인식. None이면 fmt 사용.
    """
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)

    r1 = p.add_run('- ')
    set_font(r1, 10)

    # 페이지 번호 필드. 생성 직후에는 Word 계산을 유도하되, 최종 저장 전 dirty 플래그는 제거한다.
    rpr = f'<w:rPr {nsdecls("w")}><w:sz w:val="20"/></w:rPr>'
    field_parts = [
        f'<w:r {nsdecls("w")}>{rpr}<w:fldChar w:fldCharType="begin" w:dirty="true"/></w:r>',
        f'<w:r {nsdecls("w")}>{rpr}<w:instrText xml:space="preserve"> PAGE \\* {fmt} </w:instrText></w:r>',
        f'<w:r {nsdecls("w")}>{rpr}<w:fldChar w:fldCharType="separate"/></w:r>',
        f'<w:r {nsdecls("w")}>{rpr}<w:t>1</w:t></w:r>',
        f'<w:r {nsdecls("w")}>{rpr}<w:fldChar w:fldCharType="end"/></w:r>',
    ]
    for part in field_parts:
        p._element.append(parse_xml(part))

    r2 = p.add_run(' -')
    set_font(r2, 10)

    # 시작 번호 지정 (front matter는 i부터, body는 1부터 재시작)
    # NOTE: LibreOffice/Word는 두 번째 섹션의 pgNumType.start reset이 sectPr 안에
    #       pgNumType element가 2개 이상 있을 때만 인식하는 동작이 있다. 따라서
    #       기존 element 모두 제거 + placeholder + reset 두 개를 추가한다.
    if start_num is not None:
        sectPr = section._sectPr
        for ex in list(sectPr.findall(qn('w:pgNumType'))):
            sectPr.remove(ex)
        ph_fmt = placeholder_fmt or ('lowerRoman' if fmt == 'decimal' else 'decimal')
        placeholder = parse_xml(
            f'<w:pgNumType {nsdecls("w")} w:fmt="{ph_fmt}"/>'
        )
        new_el = parse_xml(
            f'<w:pgNumType {nsdecls("w")} w:start="{start_num}" w:fmt="{fmt}"/>'
        )
        cols = sectPr.find(qn('w:cols'))
        if cols is not None:
            cols.addprevious(placeholder)
            cols.addprevious(new_el)
        else:
            sectPr.append(placeholder)
            sectPr.append(new_el)



_BM_ID = [1000]


def _next_bm_id():
    _BM_ID[0] += 1
    return _BM_ID[0]


def wrap_paragraph_with_bookmark(p, name):
    """기존 paragraph p의 콘텐츠를 bookmarkStart/End로 감싼다."""
    bid = _next_bm_id()
    bs = parse_xml(f'<w:bookmarkStart {nsdecls("w")} w:id="{bid}" w:name="{name}"/>')
    be = parse_xml(f'<w:bookmarkEnd {nsdecls("w")} w:id="{bid}"/>')
    pPr = p._p.find(qn('w:pPr'))
    if pPr is not None:
        pPr.addnext(bs)
    else:
        p._p.insert(0, bs)
    p._p.append(be)


def append_pageref_field(p, name, font=None, size=11, bold=False):
    if font is None:
        font = FONT_KR
    """Paragraph p 끝에 PAGEREF <name> \\h 필드를 삽입.
    PAGE 자리 표시는 0으로 두고 settings.xml의 updateFields가 채운다."""
    rPr_xml = f'<w:rPr><w:rFonts {nsdecls("w")} w:eastAsia="{font}"/><w:sz w:val="{size*2}"/>{("<w:b/>") if bold else ""}</w:rPr>'
    begin = parse_xml(f'<w:r {nsdecls("w")}>{rPr_xml}<w:fldChar w:fldCharType="begin" w:dirty="true"/></w:r>')
    instr = parse_xml(
        f'<w:r {nsdecls("w")}>{rPr_xml}<w:instrText xml:space="preserve"> PAGEREF {name} \\h </w:instrText></w:r>'
    )
    sep = parse_xml(f'<w:r {nsdecls("w")}>{rPr_xml}<w:fldChar w:fldCharType="separate"/></w:r>')
    placeholder = parse_xml(f'<w:r {nsdecls("w")}>{rPr_xml}<w:t>0</w:t></w:r>')
    end = parse_xml(f'<w:r {nsdecls("w")}>{rPr_xml}<w:fldChar w:fldCharType="end"/></w:r>')
    for el in (begin, instr, sep, placeholder, end):
        p._p.append(el)


def add_toc_entry(doc, text, bm_name, level=0, bold=False, tab_pos_cm=14.0, font_size=11, line_spacing=1.6):
    """점 리더 + 우측 정렬 페이지번호로 구성된 목차 항목."""
    from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(0.7 * level)
    p.paragraph_format.line_spacing = line_spacing
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.tab_stops.add_tab_stop(
        Cm(tab_pos_cm), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS
    )
    r = p.add_run(text)
    set_font(r, font_size, bold)
    rt = p.add_run('\t')
    set_font(rt, font_size, bold)
    if bm_name:
        append_pageref_field(p, bm_name, size=font_size, bold=bold)
    return p


def _empty(doc, n=1):
    """표지·제출서용 빈 줄.

    학위논문 양식의 표지 여백은 빈 줄 리듬으로 만들어지므로 과도하게 압축하지 않는다.
    """
    from docx.shared import Pt as _Pt
    for _ in range(n):
        p = doc.add_paragraph('')
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_before = _Pt(0)
        p.paragraph_format.space_after = _Pt(0)

def create_doc():
    doc = Document()

    # 기본 섹션 설정 (표지/제출서/인준서 — 세로 중앙 정렬)
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Cm(4.0)
    section.bottom_margin = Cm(4.0)
    section.left_margin = Cm(3.5)
    section.right_margin = Cm(3.5)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.5)
    # 표지·제출서·인준서 세로 가운데 정렬 (Word/한컴 표준 지원). LibreOffice는 무시.
    sectPr = section._sectPr
    for ex in list(sectPr.findall(qn('w:vAlign'))):
        sectPr.remove(ex)
    sectPr.append(parse_xml(f'<w:vAlign {nsdecls("w")} w:val="center"/>'))
    
    # 기본 스타일
    style = doc.styles['Normal']
    style.font.name = FONT_KR
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    style.paragraph_format.line_spacing = 1.6
    style.paragraph_format.first_line_indent = Cm(0.7)  # 가이드 "2칸" 들여쓰기
    rPr = style.element.get_or_add_rPr()
    rFonts = parse_xml(
        f'<w:rFonts {nsdecls("w")} '
        f'w:ascii="{FONT_KR}" w:hAnsi="{FONT_KR}" w:eastAsia="{FONT_KR}" w:cs="{FONT_KR}"/>'
    )
    rPr.append(rFonts)

    # 각주 스타일 — 가이드: 9pt 명조계열
    for fn_style_name in ('Footnote Text', 'Footnote Reference'):
        try:
            fs = doc.styles[fn_style_name]
            fs.font.name = FONT_KR
            fs.font.size = Pt(9)
            fs_rPr = fs.element.get_or_add_rPr()
            fs_rFonts = parse_xml(
                f'<w:rFonts {nsdecls("w")} '
                f'w:ascii="{FONT_KR}" w:hAnsi="{FONT_KR}" w:eastAsia="{FONT_KR}" w:cs="{FONT_KR}"/>'
            )
            fs_rPr.append(fs_rFonts)
        except KeyError:
            pass

    # Heading 1~9 디폴트 색상·크기 가이드 부합 reset.
    # 가이드: 큰제목 16pt(H1)·중간제목 13pt(H2)·본문 11pt(H3~), 모두 검정.
    heading_sizes = {1: 32, 2: 26, 3: 22, 4: 22}  # half-points
    for lvl in range(1, 10):
        try:
            hs = doc.styles[f'Heading {lvl}']
            hs_rPr = hs.element.get_or_add_rPr()
            # 색상 검정 강제
            for ex in list(hs_rPr.findall(qn('w:color'))):
                hs_rPr.remove(ex)
            hs_rPr.append(parse_xml(f'<w:color {nsdecls("w")} w:val="000000"/>'))
            # 크기 (지정된 레벨만)
            if lvl in heading_sizes:
                for ex in list(hs_rPr.findall(qn('w:sz'))):
                    hs_rPr.remove(ex)
                for ex in list(hs_rPr.findall(qn('w:szCs'))):
                    hs_rPr.remove(ex)
                hs_rPr.append(parse_xml(f'<w:sz {nsdecls("w")} w:val="{heading_sizes[lvl]}"/>'))
                hs_rPr.append(parse_xml(f'<w:szCs {nsdecls("w")} w:val="{heading_sizes[lvl]}"/>'))
        except KeyError:
            pass

    return doc


def latex_to_omml(latex, block=False):
    """LaTeX -> MathML -> OMML XML element (with m: namespace). block=True wraps in oMathPara."""
    mml = _l2m.convert(latex)
    omml_str = _m2o.convert(mml)  # starts with <m:oMath>...</m:oMath>
    # Bug workaround: mathml2omml 0.0.2 emits malformed closing tag for groupChrPr.
    omml_str = re.sub(
        r'(<m:groupChrPr>[^<]*<m:chr[^/]*/>[^<]*<m:pos[^/]*/>)\s*</m:groupChr>',
        r'\1</m:groupChrPr>',
        omml_str,
    )
    # Inject namespace on root element so parse_xml can resolve m: prefix.
    omml_ns = omml_str.replace('<m:oMath>', f'<m:oMath xmlns:m="{OMML_NS}">', 1)
    if block:
        omml_ns = (
            f'<m:oMathPara xmlns:m="{OMML_NS}">'
            f'{omml_ns}'
            f'</m:oMathPara>'
        )
    return parse_xml(omml_ns)


def add_block_math(doc, latex):
    # 수식 문단 자체만 분리되지 않게 한다. 앞뒤 본문까지 묶으면 Word에서 큰 빈 공간이 생길 수 있다.
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_together = True
    p.paragraph_format.keep_with_next = False
    try:
        elem = latex_to_omml(latex, block=True)
        p._p.append(elem)
    except Exception as e:
        r = p.add_run(latex)
        r.font.size = Pt(10)
        r.font.name = 'Cambria Math'
        r.italic = True
        print(f'  [math fallback-block] {latex[:60]} -> {e}')
    return p


def add_inline_math_run(paragraph, latex):
    try:
        elem = latex_to_omml(latex, block=False)
        paragraph._p.append(elem)
        return None
    except Exception as e:
        r = paragraph.add_run(latex)
        r.font.name = 'Cambria Math'
        r.italic = True
        print(f'  [math fallback-inline] {latex[:60]} -> {e}')
        return r


def centered_text(doc, text, size=14, bold=False, space_before=0, space_after=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.keep_together = True
    # Normal 디폴트 space_after 누적 방지
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    set_font(r, size, bold)
    return p


def add_english_title_block(doc, size=12, space_before=0):
    """표지·제출서 영문 제목을 의미 단위로 고정 줄바꿈한다."""
    lines = [
        'An Analysis of Apartment Sale Price Structure in Seoul',
        'Focusing on Distance-Based Accessibility',
        'and Spatiotemporal Heterogeneity',
    ]
    for idx, line in enumerate(lines):
        centered_text(
            doc,
            line,
            size,
            True,
            space_before=space_before if idx == 0 else 0,
        )


def add_cover(doc):
    """양식1: 표지 — 한양대 부동산융합대학원 표준 (조민지 2023 양식)
    순서: 석사학위청구논문 → 한글제목 → 영문제목 → 이름 → 학교/대학원 → 연월
    전공명은 표지에 미표기, 제출서에만 표기."""
    centered_text(doc, '석 사 학 위 청 구 논 문', 16, True, space_before=20)
    centered_text(doc, '서울시 아파트 매매가격 구조 분석', 18, True, space_before=90)
    centered_text(doc, '- 거리 기반 접근성과 시공간 이질성을 중심으로 -', 14, True, space_before=2)
    add_english_title_block(doc, size=12, space_before=10)
    # 참고 표지처럼 제목군은 위쪽에 두고 하단 정보와의 세로 간격을 넓게 확보한다.
    centered_text(doc, '박  현  근', 16, True, space_before=64)
    centered_text(doc, '한 양 대 학 교  부 동 산 융 합 대 학 원', 14, True, space_before=74)
    centered_text(doc, '2026 년  8 월', 14, True, space_before=58)
    doc.add_page_break()


def add_submission(doc):
    """양식2: 제출서.

    제출서는 표지보다 항목이 많으므로 Word 렌더링에서 이름이 다음 쪽으로 밀리지 않게
    제목 블록만 표지보다 한 단계 작게 둔다.
    """
    centered_text(doc, '석 사 학 위 청 구 논 문', 16, True)
    _empty(doc, 1)
    centered_text(doc, '서울시 아파트 매매가격 구조 분석', 17, True)
    centered_text(doc, '- 거리 기반 접근성과 시공간 이질성을 중심으로 -', 13, True)
    add_english_title_block(doc, size=12, space_before=4)
    centered_text(doc, '지도교수  고  준  호', 14, True, space_before=26)
    centered_text(doc, '이 논문을 공학 석사학위청구논문으로', 12, space_before=30)
    centered_text(doc, '제출합니다.', 12)
    centered_text(doc, '2026 년  8 월', 14, True, space_before=30)
    centered_text(doc, '한 양 대 학 교  부 동 산 융 합 대 학 원', 14, True, space_before=22)
    centered_text(doc, '도 시 · 부 동 산 빅 데 이 터 전 공', 13)
    centered_text(doc, '박  현  근', 16, True, space_before=22)
    doc.add_page_break()


def add_approval(doc):
    """양식3: 인준서 — vAlign=center + 최소 padding"""
    _empty(doc, 2)
    centered_text(doc, '이 논문을 박현근의', 14)
    centered_text(doc, '석사학위 논문으로 인준함.', 14)
    _empty(doc, 2)
    centered_text(doc, '2026 년  8 월', 14, True)
    _empty(doc, 4)
    
    committee = [
        ('심 사 위 원 장', '조  미  정'),
        ('심  사  위  원', '엄  선  용'),
        ('심  사  위  원', '고  준  호'),
    ]
    for role, name in committee:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_after = Pt(24)
        r = p.add_run(f'{role} :  {name}  (인)')
        set_font(r, 14)
    
    _empty(doc, 3)
    centered_text(doc, '한 양 대 학 교  부 동 산 융 합 대 학 원', 14, True)
    # 다음에 새 섹션이 NEW_PAGE로 시작하므로 별도 page break 불필요


def start_unnumbered_section(doc):
    """쪽번호가 없어야 하는 후면 양식(연구윤리서약서)용 새 섹션."""
    from docx.enum.section import WD_SECTION

    sec = doc.add_section(WD_SECTION.NEW_PAGE)
    sec.page_width = Mm(210); sec.page_height = Mm(297)
    sec.top_margin = Cm(4.0); sec.bottom_margin = Cm(4.0)
    sec.left_margin = Cm(3.5); sec.right_margin = Cm(3.5)
    sec.header_distance = Cm(1.5); sec.footer_distance = Cm(1.5)

    sectPr = sec._sectPr
    if sectPr.find(qn('w:type')) is None:
        type_el = parse_xml(f'<w:type {nsdecls("w")} w:val="nextPage"/>')
        pgSz = sectPr.find(qn('w:pgSz'))
        if pgSz is not None:
            pgSz.addprevious(type_el)
    for tag in ('w:vAlign', 'w:pgNumType'):
        for ex in list(sectPr.findall(qn(tag))):
            sectPr.remove(ex)
    sectPr.append(parse_xml(f'<w:vAlign {nsdecls("w")} w:val="top"/>'))

    for hf in (sec.header, sec.footer):
        hf.is_linked_to_previous = False
        for p in list(hf.paragraphs):
            p._element.getparent().remove(p._element)
    return sec


def ethics_paragraph(doc, text, size=11, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT,
                     space_before=0, line_spacing=1.35, first_indent=False,
                     font=FONT_KR):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.first_line_indent = Cm(0.7 if first_indent else 0)
    p.paragraph_format.left_indent = Cm(0)
    p.paragraph_format.right_indent = Cm(0)
    p.paragraph_format.line_spacing = line_spacing
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    set_font(r, size=size, bold=bold, font=font)
    return p


def add_research_ethics_pledges(doc):
    """공식 학위논문 양식: 국문 연구윤리서약서 → 영문 연구윤리서약서, 쪽번호 없음."""
    start_unnumbered_section(doc)

    _empty(doc, 1)
    ethics_paragraph(doc, '연구 윤리 서약서', 16, True, WD_ALIGN_PARAGRAPH.CENTER, space_before=0)
    ethics_paragraph(
        doc,
        '본인은 한양대학교 대학원생으로서 이 학위논문 작성과정에서\n'
        '다음과 같이 연구 윤리의 기본 원칙을 준수하였음을 서약합니다.',
        11,
        space_before=32,
        line_spacing=1.45,
    )
    ethics_paragraph(
        doc,
        '첫째, 지도교수의 지도를 받아 정직하고 엄정한 연구를 수행하여\n'
        '학위논문을 작성한다.',
        11,
        space_before=20,
        line_spacing=1.45,
    )
    ethics_paragraph(
        doc,
        '둘째, 논문 작성시 위조, 변조, 표절 등 학문적 진실성을 훼손하는\n'
        '어떤 연구 부정행위도 하지 않는다.',
        11,
        space_before=20,
        line_spacing=1.45,
    )
    ethics_paragraph(
        doc,
        '셋째, 논문 작성시 논문유사도 검증시스템 "카피킬러"등을 거쳐야\n'
        '한다.',
        11,
        space_before=20,
        line_spacing=1.45,
    )
    ethics_paragraph(doc, '2026년 6월 5일', 12, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=46)
    for text in (
        '학 위 명 : 공학석사',
        '학    과 : 도시·부동산빅데이터전공',
        '지도교수 : 고  준  호',
        '성    명 : 박  현  근  (서명)',
    ):
        p = ethics_paragraph(doc, text, 11, space_before=7)
        p.paragraph_format.left_indent = Cm(5.1)
    ethics_paragraph(
        doc,
        '한 양 대 학 교  부 동 산 융 합 대 학 원 장  귀 하',
        14,
        True,
        WD_ALIGN_PARAGRAPH.CENTER,
        space_before=34,
    )

    doc.add_page_break()

    _empty(doc, 1)
    ethics_paragraph(
        doc,
        'Declaration of Ethical Conduct in Research',
        15,
        True,
        WD_ALIGN_PARAGRAPH.LEFT,
        font='Times New Roman',
    )
    english_blocks = [
        'I, as a graduate student of Hanyang University, hereby declare\n'
        'that I have abided by the following Code of Research Ethics\n'
        'while writing this dissertation thesis, during my degree\n'
        'program.',
        '"First, I have strived to be honest in my conduct, to produce\n'
        'valid and reliable research conforming with the guidance of my\n'
        'thesis supervisor, and I affirm that my thesis contains honest,\n'
        'fair and reasonable conclusions based on my own careful\n'
        'research under the guidance of my thesis supervisor.',
        'Second, I have not committed any acts that may discredit or\n'
        'damage the credibility of my research. These include, but are\n'
        'not limited to: falsification, distortion of research findings or\n'
        'plagiarism.',
        'Third, I need to go through with Copykiller Program(Internet-\n'
        'based Plagiarism-prevention service) before submitting a\n'
        'thesis."',
    ]
    for idx, text in enumerate(english_blocks):
        ethics_paragraph(
            doc,
            text,
            10.5,
            font='Times New Roman',
            space_before=24 if idx == 0 else 17,
            line_spacing=1.18,
        )
    ethics_paragraph(
        doc,
        'June 5, 2026',
        12,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        font='Times New Roman',
        space_before=28,
    )
    for text in (
        'Degree : Master of Engineering',
        'Department : Urban and Real Estate Big Data Major',
        'Thesis Supervisor : Ko, Joon Ho',
        'Name : Park, Hyun Geun  (Signature)',
    ):
        p = ethics_paragraph(doc, text, 10.5, font='Times New Roman', space_before=6)
        p.paragraph_format.left_indent = Cm(2.0)


def start_front_matter_section(doc):
    """인준서 다음, 목차/표목차/그림목차/국문요지 영역 — 로마자 i부터 페이지번호."""
    from docx.enum.section import WD_SECTION
    sec = doc.add_section(WD_SECTION.NEW_PAGE)
    sec.page_width = Mm(210); sec.page_height = Mm(297)
    sec.top_margin = Cm(4.0); sec.bottom_margin = Cm(4.0)
    sec.left_margin = Cm(3.5); sec.right_margin = Cm(3.5)
    sec.header_distance = Cm(1.5); sec.footer_distance = Cm(1.5)
    sectPr = sec._sectPr
    if sectPr.find(qn('w:type')) is None:
        type_el = parse_xml(f'<w:type {nsdecls("w")} w:val="nextPage"/>')
        pgSz = sectPr.find(qn('w:pgSz'))
        if pgSz is not None:
            pgSz.addprevious(type_el)
    # 본문/front matter는 세로 상단 정렬 (vAlign=top). 표지(section 0) vAlign=center 상속 방지.
    for ex in list(sectPr.findall(qn('w:vAlign'))):
        sectPr.remove(ex)
    sectPr.append(parse_xml(f'<w:vAlign {nsdecls("w")} w:val="top"/>'))
    add_page_number(sec, start_num=1, fmt='lowerRoman')
    return sec


def add_heading(doc, text, level=1, bm_name=None):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    # 한양대 표준: 장·절은 0, 하위 제목은 참고논문 이미지 기준으로 들여쓴다.
    p.paragraph_format.line_spacing = 1.6
    p.paragraph_format.keep_with_next = True
    # 표준 Heading 스타일 적용 (TOC·목차 인식 + grep false positive 제거)
    style_name = f'Heading {level}'
    try:
        p.style = doc.styles[style_name]
    except Exception:
        pass

    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(24)
        p.paragraph_format.space_after = Pt(18)
        r = p.add_run(text)
        set_font(r, 16, True)
    elif level == 2:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(12)
        r = p.add_run(text)
        set_font(r, 13, True)
    elif level == 3:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent = Cm(0.7)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(8)
        r = p.add_run(text)
        set_font(r, 11, True)
    elif level == 4:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent = Cm(0.7)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(text)
        set_font(r, 12, True)
    if bm_name:
        wrap_paragraph_with_bookmark(p, bm_name)
    return p


def add_body(doc, text):
    if not text.strip():
        return
    is_table_note = text.strip().startswith('**주 ')
    has_long_identifier = bool(re.search(r'[A-Za-z0-9_./+\-]{18,}', text))
    font_size = 10 if is_table_note else 11
    p = doc.add_paragraph()
    # 본문은 양쪽정렬이 원칙이나, 표 주석/긴 코드성 식별자는 글자 간격 벌어짐을 막기 위해 예외 처리한다.
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT if (is_table_note or has_long_identifier) else WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.35 if is_table_note else 1.6
    p.paragraph_format.first_line_indent = Cm(0.5 if is_table_note else 0.7)
    p.paragraph_format.space_after = Pt(0)
    # 한·영 혼용 양쪽 정렬 시 글자 spread 방지 (LibreOffice/Word 동아시아 spacing 제어)
    pPr = p._p.get_or_add_pPr()
    for tag in ('w:autoSpaceDE', 'w:autoSpaceDN', 'w:adjustRightInd'):
        for ex in list(pPr.findall(qn(tag))):
            pPr.remove(ex)
        pPr.append(parse_xml(f'<{tag} {nsdecls("w")} w:val="0"/>'))
    # split by bold first
    bold_parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for bp in bold_parts:
        is_bold = bp.startswith('**') and bp.endswith('**')
        segment = bp[2:-2] if is_bold else bp
        # split by inline math $...$
        math_parts = re.split(r'(\$[^$\n]+\$)', segment)
        for mp in math_parts:
            if mp.startswith('$') and mp.endswith('$') and len(mp) > 2:
                add_inline_math_run(p, mp[1:-1])
            elif mp:
                # 인라인 코드 `...` 처리 — monospace 표시
                code_parts = re.split(r'(`[^`\n]+`)', mp)
                for cp in code_parts:
                    if cp.startswith('`') and cp.endswith('`') and len(cp) > 2:
                        r = p.add_run(cp[1:-1])
                        set_font(r, font_size, bold=is_bold)
                    elif cp:
                        # 단일 별표 *italic* 처리 (참고문헌 학술지명·책명 등)
                        italic_parts = re.split(r'(\*[^*\n]+\*)', cp)
                        for ip in italic_parts:
                            if ip.startswith('*') and ip.endswith('*') and len(ip) > 2 and not ip.startswith('**'):
                                r = p.add_run(ip[1:-1])
                                set_font(r, font_size, bold=is_bold, italic=True)
                            elif ip:
                                r = p.add_run(ip)
                                set_font(r, font_size, is_bold)
    return p


def is_sep(line):
    s = line.strip()
    if not (s.startswith('|') and s.endswith('|')):
        return False
    # 컬럼별로 split 후 각 칸이 정렬 마커(- : 공백)만 포함하는지
    cells = [c.strip() for c in s.strip('|').split('|')]
    if not cells:
        return False
    return all(re.fullmatch(r':?-+:?', c) for c in cells if c != '')


def parse_row(line):
    return [c.strip() for c in line.strip().strip('|').split('|')]


def add_table(doc, rows, caption='', bm_name=None):
    is_var_table = caption.startswith('<표 3-1>')
    is_distance_table = caption.startswith('<표 4-2>')
    is_ablation_design_table = caption.startswith('<표 4-6>')
    is_year_top5_overall_table = caption.startswith('<표 4-12>')
    is_year_top3_table = caption.startswith('<표 4-14>')
    is_year_top5_table = caption.startswith('<표 4-15>')
    if is_var_table:
        doc.add_page_break()
    if caption:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(caption)
        set_font(r, 10, False, FONT_KR)  # 가이드: 표/그림 캡션 10pt, 본문과 동일 폰트 패밀리
        if bm_name:
            wrap_paragraph_with_bookmark(p, bm_name)
    
    if not rows:
        return
    
    ncols = max(len(r) for r in rows)
    keep_whole_table = len(rows) <= 12
    tbl = doc.add_table(rows=len(rows), cols=ncols)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    if is_var_table and ncols == 5:
        # A4 본문폭(약 14cm)에 맞춘 변수 정의표 전용 폭.
        # 셀 내부 줄간격을 별도 압축해 불필요한 행 높이를 줄인다.
        set_table_grid(tbl, [1.7, 2.55, 5.05, 1.5, 3.2])
    elif is_distance_table and ncols == 5:
        # 거리 변수명은 첫 열에서 줄바꿈이 생기기 쉬우므로 첫 열을 넓힌다.
        set_table_grid(tbl, [4.1, 2.45, 2.45, 2.45, 2.5])
    elif is_ablation_design_table and ncols == 3:
        set_table_grid(tbl, [2.6, 6.0, 5.3])
    elif is_year_top3_table and ncols == 4:
        set_table_grid(tbl, [1.5, 4.2, 4.2, 4.1])
    elif is_year_top5_overall_table and ncols == 6:
        # 긴 변수명이 들어가는 전체 연도별 상위 5개 표는 한 페이지 안에서
        # 셀 내 줄바꿈이 과도해지지 않도록 연도 칸을 줄이고 본문 칸을 넓힌다.
        set_table_grid(tbl, [1.15, 2.57, 2.57, 2.57, 2.57, 2.57])
    elif is_year_top5_table and ncols == 6:
        # 연도 칸이 2025처럼 네 자리 숫자를 줄바꿈하지 않도록 고정 폭을 둔다.
        set_table_grid(tbl, [1.5, 2.5, 2.5, 2.5, 2.5, 2.5])
    
    for i, row_data in enumerate(rows):
        trPr = tbl.rows[i]._tr.get_or_add_trPr()
        if trPr.find(qn('w:cantSplit')) is None:
            trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))
        if i == 0 and trPr.find(qn('w:tblHeader')) is None:
            trPr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))
        for j, txt in enumerate(row_data):
            if j >= ncols:
                continue
            cell = tbl.rows[i].cells[j]
            cell.text = ''
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if is_var_table or is_distance_table or is_ablation_design_table or is_year_top5_overall_table:
                set_cell_margins(cell, top=45, left=70, bottom=45, right=70)
            else:
                set_cell_margins(cell, top=60, left=90, bottom=60, right=90)
            p = cell.paragraphs[0]
            if (is_var_table and i > 0 and j in (1, 2, 4)) or (is_ablation_design_table and i > 0 and j in (1, 2)):
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.keep_together = True
            if keep_whole_table and i < len(rows) - 1:
                p.paragraph_format.keep_with_next = True
            # 셀 내 마크다운 **bold**, *italic*, `code` 처리
            parts = re.split(r'(\*\*[^*]+\*\*|\*[^*\n]+\*|`[^`]+`)', txt)
            cell_font_size = 8.0 if is_var_table else 9
            if is_year_top3_table:
                cell_font_size = 8.6
            elif is_ablation_design_table:
                cell_font_size = 8.1
            elif is_year_top5_overall_table:
                cell_font_size = 7.5
            elif is_year_top5_table:
                cell_font_size = 8.1
            for part in parts:
                if not part:
                    continue
                if part.startswith('**') and part.endswith('**'):
                    r = p.add_run(part[2:-2])
                    set_font(r, cell_font_size, bold=True, font=FONT_KR)
                elif part.startswith('*') and part.endswith('*') and len(part) > 2 and not part.startswith('**'):
                    r = p.add_run(part[1:-1])
                    set_font(r, cell_font_size, bold=False, font=FONT_KR, italic=True)
                elif part.startswith('`') and part.endswith('`'):
                    r = p.add_run(part[1:-1])
                    set_font(r, cell_font_size, bold=(i == 0), font=FONT_KR)
                else:
                    r = p.add_run(part)
                    set_font(r, cell_font_size, bold=(i == 0), font=FONT_KR)
            if i == 0:
                set_shading(cell, 'D9E2F3')
    
    _empty(doc)
def add_image(doc, path, caption='', bm_name=None):
    if not os.path.exists(path):
        add_body(doc, f'[그림 파일 없음: {os.path.basename(path)}]')
        return
    # 이미지를 먼저 배치하고 캡션은 아래 가운데에 둔다.
    compact = os.path.basename(path) == 'fig1_research_flow.png'
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(0 if compact else 6)
    p.paragraph_format.space_after = Pt(0 if compact else 6)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.keep_together = True
    r = p.add_run()
    r.add_picture(path, width=Cm(image_width_cm_for_docx(path)))
    if caption:
        p.add_run().add_break()
        r = p.add_run(caption)
        set_font(r, 10, font=FONT_KR)
        if bm_name:
            wrap_paragraph_with_bookmark(p, bm_name)


_TOC_HEADINGS = ('목차', '표 목차', '그림 목차')


def _norm_ref_text(text):
    """목차 원문과 본문 제목의 공백 차이를 흡수하는 참조 키."""
    return re.sub(r'\s+', '', text or '')


def _heading_bookmark(bm_map, text):
    return (
        bm_map.get('headings', {}).get(text)
        or bm_map.get('headings_norm', {}).get(_norm_ref_text(text))
    )


def prebuild_bookmark_map(md_path):
    """본문 헤딩(H1·H2)·표 캡션·그림 캡션에 부여할 bookmark 사전.

    목차/표목차/그림목차 영역(H1) 안의 H2 라인은 등록하지 않는다.
    """
    headings = {}  # full text → bm name
    headings_norm = {}  # whitespace-normalized text → bm name
    tables = {}    # 'N-M' → bm name
    figures = {}   # 'N-M' → bm name

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_toc_zone = False
    seq = 0
    for ln in lines:
        s = ln.rstrip()
        if s.startswith('# ') and not s.startswith('## '):
            txt = s[2:].strip()
            seq += 1
            bm = f'bm_h_{seq}'
            headings[txt] = bm
            headings_norm[_norm_ref_text(txt)] = bm
            in_toc_zone = txt in _TOC_HEADINGS
            continue
        if s.startswith('## '):
            if in_toc_zone:
                continue
            txt = s[3:].strip()
            seq += 1
            bm = f'bm_h_{seq}'
            headings[txt] = bm
            headings_norm[_norm_ref_text(txt)] = bm
            continue
        m = re.match(r'^(?:\*\*)?<표\s*([0-9-]+)>.*?(?:\*\*)?$', s.strip())
        if m:
            num = m.group(1)
            tables[num] = f'bm_t_{num.replace("-", "_")}'
            continue
        m = re.match(r'^!\[<그림\s*([0-9-]+)>[^]]*\]\(.+?\)\s*$', s.strip())
        if m:
            num = m.group(1)
            figures[num] = f'bm_f_{num.replace("-", "_")}'
            continue
    return {'headings': headings, 'headings_norm': headings_norm, 'tables': tables, 'figures': figures}


def convert_md(doc, md_path, bm_map=None):
    if bm_map is None:
        bm_map = {'headings': {}, 'headings_norm': {}, 'tables': {}, 'figures': {}}

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    # 본문 시작은 표제지/제출서/인준서 다음의 첫 marker(목차 또는 국문요지)부터
    start = 0
    for i, line in enumerate(lines):
        s = line.strip()
        if (
            s.startswith('# 목차')
            or s.startswith('# 국문요지')
            or s.startswith('# 국문 요지')
            or s.startswith('# 국문초록')
            or s.startswith('# 국문 초록')
        ):
            start = i
            break

    i = start
    caption = ''
    body_section_started = False  # 첫 '제N장' 헤딩을 만나면 새 섹션 시작
    just_started_section = True   # 직전이 섹션 break이면 H1에서 page break 중복 방지 (front matter 진입 직후 True)
    in_references = False  # 참고문헌 섹션 hanging indent 모드
    toc_mode = None  # 'chapter' | 'table' | 'figure' | None

    while i < len(lines):
        line = lines[i].rstrip()
        
        # 블록 수식 (한 줄 $$...$$)
        m_one = re.match(r'^\s*\$\$(.+)\$\$\s*$', line)
        if m_one:
            add_block_math(doc, m_one.group(1).strip())
            i += 1
            continue
        # 블록 수식 (다중 라인 $$ ... $$)
        if line.strip() == '$$':
            parts = []
            i += 1
            while i < len(lines) and lines[i].strip() != '$$':
                parts.append(lines[i].strip())
                i += 1
            add_block_math(doc, ' '.join(parts))
            i += 1
            continue
        
        # 표 캡션
        m = re.match(r'^(?:\*\*)?(<표[^>]*>.*?)(?:\*\*)?$', line.strip())
        if m:
            caption = m.group(1)
            i += 1
            continue

        # 표
        if '|' in line and line.strip().startswith('|') and not is_sep(line):
            rows = []
            while i < len(lines):
                l = lines[i].rstrip()
                if '|' in l and l.strip().startswith('|'):
                    if not is_sep(l):
                        rows.append(parse_row(l))
                    i += 1
                else:
                    break
            tbm = None
            mn = re.match(r'^<표\s*([0-9-]+)>', caption)
            if mn:
                tbm = bm_map['tables'].get(mn.group(1))
            add_table(doc, rows, caption, bm_name=tbm)
            caption = ''
            continue

        if is_sep(line):
            i += 1
            continue

        # 그림
        m = re.match(r'!\[(.+?)\]\((.+?)\)', line)
        if m:
            cap = m.group(1)
            rel = m.group(2)
            absp = os.path.normpath(os.path.join(PAPER_DIR, rel))
            fbm = None
            mn = re.match(r'^<그림\s*([0-9-]+)>', cap)
            if mn:
                fbm = bm_map['figures'].get(mn.group(1))
            add_image(doc, absp, cap, bm_name=fbm)
            i += 1
            continue

        # 제목
        if line.startswith('# '):
            txt = line[2:].strip()
            # 참고문헌/Abstract 헤딩에서 모드 전환
            if txt.startswith('참고문헌'):
                in_references = True
            elif txt.startswith('Abstract') or txt.startswith('ABSTRACT'):
                in_references = False
            # 목차/표목차/그림목차 영역 진입 플래그
            if txt == '목차':
                toc_mode = 'chapter'
            elif txt == '표 목차':
                toc_mode = 'table'
            elif txt == '그림 목차':
                toc_mode = 'figure'
            else:
                toc_mode = None
            # 장 제목은 새 페이지. 첫 제N장 만남 → body 섹션 시작 (페이지 번호 1부터 아라비아)
            is_chapter = txt.startswith('제') and '장' in txt
            needs_page_break_before = False
            if is_chapter and not body_section_started:
                # 첫 본문 chapter (제1장) — 새 섹션 (자동 페이지 break)
                from docx.enum.section import WD_SECTION
                new_sec = doc.add_section(WD_SECTION.NEW_PAGE)
                new_sec.page_width = Mm(210)
                new_sec.page_height = Mm(297)
                new_sec.top_margin = Cm(4.0)
                new_sec.bottom_margin = Cm(4.0)
                new_sec.left_margin = Cm(3.5)
                new_sec.right_margin = Cm(3.5)
                new_sec.header_distance = Cm(1.5)
                new_sec.footer_distance = Cm(1.5)
                sectPr = new_sec._sectPr
                if sectPr.find(qn('w:type')) is None:
                    type_el = parse_xml(f'<w:type {nsdecls("w")} w:val="nextPage"/>')
                    pgSz = sectPr.find(qn('w:pgSz'))
                    if pgSz is not None:
                        pgSz.addprevious(type_el)
                    else:
                        sectPr.append(type_el)
                for ex in list(sectPr.findall(qn('w:vAlign'))):
                    sectPr.remove(ex)
                sectPr.append(parse_xml(f'<w:vAlign {nsdecls("w")} w:val="top"/>'))
                add_page_number(new_sec, start_num=1, fmt='decimal')
                body_section_started = True
                just_started_section = False  # 제1장 헤딩 = 본문 새 섹션 첫 paragraph → 자체 break 불요
            elif is_chapter:
                # 본문 후속 chapter (제2~5장) — 명시적 page break
                needs_page_break_before = True
                just_started_section = False
            elif just_started_section:
                # front matter 진입 직후 첫 H1 (목차) — 섹션 break이 이미 페이지 break 효과
                just_started_section = False
            else:
                # 그 외 모든 H1 (표목차·그림목차·국문요지·참고문헌·Abstract·감사의 글)
                needs_page_break_before = True
            hbm = _heading_bookmark(bm_map, txt)
            p_h1 = add_heading(doc, txt, 1, bm_name=hbm)
            if needs_page_break_before and p_h1 is not None:
                p_h1.paragraph_format.page_break_before = True
            i += 1
            continue
        if line.startswith('## '):
            txt = line[3:].strip()
            # 챕터 목차 영역에서는 ## 제N장 ... 라인을 TOC 항목으로 변환
            if toc_mode == 'chapter':
                bm = _heading_bookmark(bm_map, txt)
                add_toc_entry(doc, txt, bm, level=0, bold=True)
                i += 1
                continue
            hbm = _heading_bookmark(bm_map, txt)
            add_heading(doc, txt, 2, bm_name=hbm)
            i += 1
            continue
        if line.startswith('#### '):
            add_heading(doc, line[5:].strip(), 4)
            i += 1
            continue
        if line.startswith('### '):
            add_heading(doc, line[4:].strip(), 3)
            i += 1
            continue
        
        if line.strip() == '---':
            # 다음 비공백 줄이 H1 헤딩이면 page break 생략 (H1 자체가 처리하므로 중복 방지)
            j = i + 1
            while j < len(lines) and not lines[j].strip():
                j += 1
            if j < len(lines) and lines[j].startswith('# '):
                i += 1
                continue
            doc.add_page_break()
            i += 1
            continue
        
        # 표 목차·그림 목차의 항목 → TOC entry (점 리더 + PAGEREF)
        if toc_mode in ('table', 'figure') and (
            line.strip().startswith('- <표') or line.strip().startswith('- <그림')
        ):
            txt = line.strip()[2:]  # "- " 제거
            mn = re.match(r'^<(표|그림)\s*([0-9-]+)>', txt)
            bm = None
            if mn:
                key = mn.group(2)
                if mn.group(1) == '표':
                    bm = bm_map['tables'].get(key)
                else:
                    bm = bm_map['figures'].get(key)
            if toc_mode in ('table', 'figure'):
                add_toc_entry(doc, txt, bm, level=0, font_size=10.5, line_spacing=1.35)
            else:
                add_toc_entry(doc, txt, bm, level=0)
            i += 1
            continue

        if not line.strip():
            i += 1
            continue

        # 마크다운 리스트 → docx bullet/number 스타일
        m_dash = re.match(r'^[ \t]*-\s+(.*)$', line)
        m_num = re.match(r'^[ \t]*(\d+)\.\s+(.*)$', line)
        if m_dash:
            txt = m_dash.group(1)
            # 챕터 TOC 영역의 bullet 항목 → TOC entry
            if toc_mode == 'chapter':
                # 절 라인은 들여쓰기 1단, 그 외(국문요지·표목차 등)는 들여쓰기 0단·굵게
                if txt.startswith('제') and '절' in txt[:6]:
                    bm = _heading_bookmark(bm_map, txt)
                    add_toc_entry(doc, txt, bm, level=1)
                else:
                    bm = _heading_bookmark(bm_map, txt)
                    add_toc_entry(doc, txt, bm, level=0, bold=True)
                i += 1
                continue
            p = add_body(doc, txt)
            if p is not None:
                if in_references:
                    # 참고문헌: bullet 제거, hanging indent (1.0cm)
                    p.paragraph_format.first_line_indent = Cm(-1.0)
                    p.paragraph_format.left_indent = Cm(1.0)
                    p.paragraph_format.widow_control = True  # 마지막 줄 widow 방지 (keep_together는 부작용)
                else:
                    p.style = doc.styles['List Bullet']
                    p.paragraph_format.first_line_indent = Cm(0)
            i += 1
            continue
        if m_num:
            p = add_body(doc, m_num.group(2))
            if p is not None:
                p.style = doc.styles['List Number']
                p.paragraph_format.first_line_indent = Cm(0)
            i += 1
            continue

        add_body(doc, line)
        i += 1


def enable_update_fields(doc):
    """settings.xml에 <w:updateFields w:val='true'/> 추가 — open 시 PAGE 등 필드 자동 재계산.
    추가로 한국어 양쪽 정렬 글자 spread 완화를 위한 character spacing 옵션도 함께 박는다."""
    settings = doc.settings.element
    # 1. updateFields
    if settings.find(qn('w:updateFields')) is None:
        settings.insert(0, parse_xml(f'<w:updateFields {nsdecls("w")} w:val="true"/>'))
    # 2. 한국어 자간 spread 완화 옵션들
    extras = [
        # 구두점 + 일본어 가나 압축 (한국어에도 적용, 가장 강한 compression)
        '<w:characterSpacingControl {ns} w:val="compressPunctuationAndJapaneseKana"/>',
        # 동아시아 단어 단위 줄바꿈
        '<w:wordWrap {ns} w:val="0"/>',
        # 구두점 overflow 허용
        '<w:overflowPunct {ns} w:val="0"/>',
        # kinsoku 규칙 완화
        '<w:strictFirstAndLastChars {ns} w:val="0"/>',
        # 양쪽 정렬 끝 공백 처리
        '<w:doNotExpandShiftReturn {ns}/>',
        # Word95 호환 spacing (한국어 spread 보수적 처리)
        '<w:autoSpaceLikeWord95 {ns}/>',
        # HTML auto-spacing 비활성
        '<w:doNotUseHTMLParagraphAutoSpacing {ns}/>',
        # 균등 분배 조정 비활성
        '<w:adjustLineHeightInTable {ns}/>',
    ]
    for x in extras:
        tag = x.split()[0][1:]  # 'w:xxx'
        if settings.find(qn(tag)) is None:
            settings.append(parse_xml(x.format(ns=nsdecls('w'))))


def freeze_fields_via_libreoffice(docx_path):
    """LibreOffice headless로 docx 재저장 → PAGE/PAGEREF 필드 평가값을 OOXML cache에 굽기.

    Word for Mac AppleScript는 macOS sandbox로 fields update를 OOXML까지 굽지 못한다
    (sdef 비어있음, do Visual Basic 컴파일 차단). LibreOffice는 변환 단계에서 fields를
    평가하고 cache를 굽기 때문에 후처리로 사용한다. 한글 폰트·표 정렬에 미세한 변동이
    있을 수 있어 결과 docx는 시각 검증 필요.
    """
    import subprocess, shutil, tempfile
    soffice = '/Applications/LibreOffice.app/Contents/MacOS/soffice'
    if not os.path.exists(soffice):
        print('  [skip] LibreOffice 미설치 — 필드는 placeholder로 남음')
        return False
    with tempfile.TemporaryDirectory() as tmp:
        try:
            subprocess.run(
                [soffice, '--headless', '--convert-to', 'docx',
                 '--outdir', tmp, docx_path],
                check=True, timeout=240,
                stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL,
            )
        except Exception as e:
            print(f'  [warn] field freeze 실패: {e}')
            return False
        fname = os.path.basename(docx_path)
        src = os.path.join(tmp, fname)
        if not os.path.exists(src):
            print('  [warn] LibreOffice 출력 파일 없음')
            return False
        shutil.move(src, docx_path)
        print(f'  ✓ PAGE/PAGEREF 필드 cache 채움 (LibreOffice 재저장)')
        return True


def export_pdf_via_libreoffice(docx_path):
    """docx → PDF (LibreOffice headless)."""
    import subprocess
    soffice = '/Applications/LibreOffice.app/Contents/MacOS/soffice'
    if not os.path.exists(soffice):
        return False
    try:
        subprocess.run(
            [soffice, '--headless', '--convert-to', 'pdf',
             '--outdir', os.path.dirname(docx_path), docx_path],
            check=True, timeout=240,
            stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL,
        )
        print(f'  ✓ PDF 변환 완료')
        return True
    except Exception as e:
        print(f'  [warn] PDF 변환 실패: {e}')
        return False


def export_pdf_via_word(docx_path, pdf_path):
    """docx → PDF (Microsoft Word). LibreOffice보다 한국어 폰트·정렬 정확."""
    import shutil
    import subprocess
    import tempfile

    docx_path = os.path.abspath(docx_path)
    pdf_path = os.path.abspath(pdf_path)
    if os.path.exists(pdf_path):
        os.remove(pdf_path)
    if os.name == 'nt':
        tmpdir = tempfile.mkdtemp(prefix='thesis_word_export_')
        ps1 = os.path.join(tmpdir, 'export_word_pdf.ps1')
        script = r'''
param(
  [Parameter(Mandatory=$true)][string]$DocxPath,
  [Parameter(Mandatory=$true)][string]$PdfPath
)
$ErrorActionPreference = "Stop"
$word = New-Object -ComObject Word.Application
$word.Visible = $false
$word.DisplayAlerts = 0
try {
  $doc = $word.Documents.Open($DocxPath, $false, $false)
  foreach ($field in $doc.Fields) { $field.Update() | Out-Null }
  foreach ($toc in $doc.TablesOfContents) { $toc.Update() | Out-Null }
  foreach ($tof in $doc.TablesOfFigures) { $tof.Update() | Out-Null }
  $doc.Save()
  $doc.ExportAsFixedFormat($PdfPath, 17, $false, 0)
  $doc.Close($false)
}
finally {
  $word.Quit()
}
'''
        try:
            with open(ps1, 'w', encoding='utf-8') as f:
                f.write(script)
            subprocess.run(
                [
                    'powershell',
                    '-NoProfile',
                    '-ExecutionPolicy',
                    'Bypass',
                    '-File',
                    ps1,
                    '-DocxPath',
                    docx_path,
                    '-PdfPath',
                    pdf_path,
                ],
                check=True,
                timeout=240,
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL,
            )
            print(f'  ✓ Word PDF 변환 완료 ({FONT_KR}, 한국어 정렬 정확)')
            return os.path.exists(pdf_path)
        except Exception as e:
            print(f'  [warn] Word PDF 변환 실패: {e}')
            return False
        finally:
            shutil.rmtree(tmpdir, ignore_errors=True)

    tmpdir = tempfile.mkdtemp(prefix='thesis_word_export_')
    word_docx = os.path.join(tmpdir, 'thesis.docx')
    word_pdf = os.path.join(tmpdir, 'thesis.pdf')
    shutil.copy2(docx_path, word_docx)
    word_was_running = subprocess.run(
        ['pgrep', '-x', 'Microsoft Word'],
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
    ).returncode == 0
    script = f'''
set theDocPath to "{word_docx}"
set thePdfPath to POSIX file "{word_pdf}"
tell application "Microsoft Word"
    launch
    delay 1
    open file name theDocPath
    delay 3
    set theDoc to active document
    save theDoc
    save as theDoc file name (thePdfPath as string) file format format PDF
end tell
'''
    try:
        proc = subprocess.run(['osascript', '-e', script], check=False, timeout=180,
                              stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
        if os.path.exists(word_pdf):
            shutil.copy2(word_pdf, pdf_path)
            if not word_was_running:
                subprocess.run(
                    ['osascript', '-e', 'tell application "Microsoft Word" to quit saving no'],
                    stdout=subprocess.DEVNULL,
                    stderr=subprocess.DEVNULL,
                    timeout=10,
                    check=False,
                )
                subprocess.run(
                    ['pkill', '-x', 'Microsoft Word'],
                    stdout=subprocess.DEVNULL,
                    stderr=subprocess.DEVNULL,
                    check=False,
                )
            if proc.returncode != 0:
                print('  [warn] Word가 종료 응답을 반환했지만 PDF 저장은 완료됨')
            print(f'  ✓ Word PDF 변환 완료 ({FONT_KR}, 한국어 정렬 정확)')
            return True
        return False
    except Exception as e:
        print(f'  [warn] Word PDF 변환 실패: {e}')
        return False
    finally:
        if not word_was_running:
            subprocess.run(
                ['osascript', '-e', 'tell application "Microsoft Word" to quit saving no'],
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL,
                timeout=10,
                check=False,
            )
            subprocess.run(
                ['pkill', '-x', 'Microsoft Word'],
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL,
                check=False,
            )
        shutil.rmtree(tmpdir, ignore_errors=True)


def _toc_static_key(text):
    """PDF 추출/Word 렌더의 공백·점리더·중점 차이를 흡수하는 TOC 비교 키."""
    text = re.sub(r'[.·•…]+', '', text or '')
    return re.sub(r'\s+', '', text)


def _pdf_footer_page_number(page_text):
    matches = re.findall(
        r'(?m)^\s*-\s*([0-9ivxlcdmIVXLCDM]+)\s*-\s*$',
        page_text or '',
    )
    return matches[-1] if matches else None


def _extract_pdf_pages(pdf_path):
    import subprocess

    try:
        raw = subprocess.check_output(
            ['pdftotext', '-layout', pdf_path, '-'],
            stderr=subprocess.DEVNULL,
        ).decode('utf-8', 'replace')
    except Exception as e:
        print(f'  [warn] PDF 텍스트 추출 실패: {e}')
        return []
    return [p for p in raw.split('\f') if p.strip()]


def _extract_toc_entries_from_pdf(pdf_path):
    """PDF의 목차/표목차/그림목차 영역에서 (항목명, 페이지 번호)를 순서대로 추출한다."""
    pages = _extract_pdf_pages(pdf_path)
    if not pages:
        return []

    first_body_idx = next(
        (idx for idx, page in enumerate(pages) if _pdf_footer_page_number(page) == '1'),
        len(pages),
    )

    toc_entries = []
    for page in pages[:first_body_idx]:
        for line in page.splitlines():
            raw_line = line.strip()
            if not raw_line or raw_line in ('목차', '표 목차', '그림 목차'):
                continue
            # 점 리더는 pdftotext에서 연속 마침표로 추출된다.
            cleaned = re.sub(r'\.{2,}', ' ', raw_line)
            m = re.match(
                r'^(?P<label>.+?)\s+(?P<num>[0-9]+|[ivxlcdmIVXLCDM]+)\s*$',
                cleaned,
            )
            if not m:
                continue
            label = m.group('label').strip()
            if not (
                label.startswith('제')
                or label.startswith('<표')
                or label.startswith('<그림')
                or label in ('참고문헌', 'Abstract')
            ):
                continue
            toc_entries.append((label, m.group('num')))
    return toc_entries


def _extract_toc_numbers_from_pdf(pdf_path):
    """Word가 계산한 목차 페이지 번호를 PDF의 목차 영역에서 추출한다."""
    toc_entries = _extract_toc_entries_from_pdf(pdf_path)
    if not toc_entries:
        print('  [warn] PDF에서 목차 번호를 추출하지 못함')
        return {}
    return {_toc_static_key(label): num for label, num in toc_entries}


def _toc_caption_id(text):
    m = re.match(r'^\s*(<\s*(?:표|그림)\s*[0-9]+(?:-[0-9]+)?\s*>)', text or '')
    return _toc_static_key(m.group(1)) if m else None


def _toc_entry_matches_pdf_line(label, line):
    """본문 참조문·그림 내부 텍스트를 실제 제목으로 오인하지 않도록 엄격히 매칭한다."""
    caption_id = _toc_caption_id(label)
    if caption_id:
        return _toc_caption_id(line.strip()) == caption_id
    return _toc_static_key(line.strip()) == _toc_static_key(label)


def _extract_actual_toc_numbers_from_pdf(pdf_path, labels):
    pages = _extract_pdf_pages(pdf_path)
    if not pages:
        return {}

    first_body_idx = next(
        (idx for idx, page in enumerate(pages) if _pdf_footer_page_number(page) == '1'),
        len(pages),
    )

    actual_numbers = {}
    for label in labels:
        label_key = _toc_static_key(label)
        for page in pages[first_body_idx:]:
            page_num = _pdf_footer_page_number(page)
            if not (page_num and page_num.isdigit()):
                continue
            for line in page.splitlines():
                if _toc_entry_matches_pdf_line(label, line):
                    actual_numbers[label_key] = page_num
                    break
            if label_key in actual_numbers:
                break
    return actual_numbers


def validate_toc_numbers_in_pdf(pdf_path):
    """최종 PDF에서 목차 표기 번호와 실제 본문/표/그림 위치를 전수 비교한다."""
    entries = _extract_toc_entries_from_pdf(pdf_path)
    if not entries:
        print('  [warn] 목차 검증 실패: PDF 목차 항목을 추출하지 못함')
        return False, {}

    actual_numbers = _extract_actual_toc_numbers_from_pdf(
        pdf_path,
        [label for label, _num in entries],
    )
    missing = []
    mismatches = []
    for label, toc_num in entries:
        label_key = _toc_static_key(label)
        actual_num = actual_numbers.get(label_key)
        if actual_num is None:
            missing.append(label)
        elif str(toc_num) != str(actual_num):
            mismatches.append((label, toc_num, actual_num))

    report = {
        'entries': entries,
        'actual_numbers': actual_numbers,
        'missing': missing,
        'mismatches': mismatches,
    }
    if missing or mismatches:
        print(
            f'  [warn] 목차 페이지 번호 검증 실패: '
            f'누락 {len(missing)}개, 불일치 {len(mismatches)}개'
        )
        for label, toc_num, actual_num in mismatches[:8]:
            print(f'    - {label}: 목차 {toc_num} / 실제 {actual_num}')
        if missing:
            print(f'    - 실제 위치 미검출: {missing[:5]}')
        return False, report

    print(f'  ✓ 목차 페이지 번호 검증 통과: {len(entries)}개 항목')
    return True, report


def _extract_toc_pageref_targets_from_docx(docx_path):
    """DOCX 목차 필드에서 (표시 라벨, bookmark 이름)을 추출한다."""
    from zipfile import ZipFile
    from lxml import etree

    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    with ZipFile(docx_path, 'r') as zin:
        document_xml = zin.read('word/document.xml')

    root = etree.fromstring(document_xml)
    pairs = []
    for p in root.xpath('//w:p[.//w:instrText[contains(., "PAGEREF")]]', namespaces=ns):
        label = _toc_label_before_pageref(p, ns)
        instr = ''.join(p.xpath('.//w:instrText/text()', namespaces=ns))
        m = re.search(r'\bPAGEREF\s+([A-Za-z0-9_]+)\b', instr)
        if label and m:
            pairs.append((label, m.group(1)))
    return pairs


def _rewrite_docx_parts(docx_path, part_updater):
    """DOCX zip 내부 part를 필요한 경우에만 갱신한다."""
    import tempfile
    from zipfile import ZipFile, ZIP_DEFLATED

    with ZipFile(docx_path, 'r') as zin:
        items = zin.infolist()
        data = {item.filename: zin.read(item.filename) for item in items}

    changed = False
    for name, original in list(data.items()):
        updated = part_updater(name, original)
        if updated is not None and updated != original:
            data[name] = updated
            changed = True

    if not changed:
        return False

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    tmp.close()
    try:
        with ZipFile(tmp.name, 'w', ZIP_DEFLATED) as zout:
            for item in items:
                zout.writestr(item, data[item.filename])
        os.replace(tmp.name, docx_path)
    finally:
        if os.path.exists(tmp.name):
            os.remove(tmp.name)
    return True


def remove_update_fields_from_docx(docx_path):
    """Word 열기 시 필드 업데이트 경고를 띄우는 updateFields 설정을 제거한다."""
    from lxml import etree

    ns_path = 'word/settings.xml'

    def updater(name, original):
        if name != ns_path:
            return None
        root = etree.fromstring(original)
        changed = False
        for el in list(root.findall(qn('w:updateFields'))):
            root.remove(el)
            changed = True
        if not changed:
            return None
        return etree.tostring(
            root,
            xml_declaration=True,
            encoding='UTF-8',
            standalone=True,
        )

    return _rewrite_docx_parts(docx_path, updater)


def clear_dirty_fields_from_docx(docx_path):
    """최종 DOCX에서 Word 필드 자동 갱신 유발 플래그를 제거한다."""
    from lxml import etree

    def updater(name, original):
        if not (name.startswith('word/') and name.endswith('.xml')):
            return None
        try:
            root = etree.fromstring(original)
        except etree.XMLSyntaxError:
            return None
        changed = False
        for el in root.xpath('//*[@w:dirty]', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
            del el.attrib[qn('w:dirty')]
            changed = True
        if not changed:
            return None
        return etree.tostring(
            root,
            xml_declaration=True,
            encoding='UTF-8',
            standalone=True,
        )

    return _rewrite_docx_parts(docx_path, updater)


def sanitize_final_docx_fields(docx_path):
    """제출 DOCX가 열릴 때 필드 업데이트 팝업으로 목차가 흔들리지 않게 정리한다."""
    removed_update_fields = remove_update_fields_from_docx(docx_path)
    cleared_dirty_fields = clear_dirty_fields_from_docx(docx_path)
    return removed_update_fields or cleared_dirty_fields


def _applescript_quote(text):
    return '"' + (text or '').replace('\\', '\\\\').replace('"', '\\"') + '"'


def _word_bookmark_page_numbers(docx_path, bookmark_names):
    """Microsoft Word가 DOCX 화면에서 계산한 bookmark별 조정 페이지 번호를 읽는다."""
    import shutil
    import subprocess
    import tempfile

    unique_names = []
    for name in bookmark_names:
        if name and name not in unique_names:
            unique_names.append(name)
    if not unique_names:
        return {}
    if os.name == 'nt':
        tmpdir = tempfile.mkdtemp(prefix='thesis_word_probe_')
        probe_docx = os.path.join(tmpdir, 'probe.docx')
        list_path = os.path.join(tmpdir, 'bookmarks.txt')
        out_path = os.path.join(tmpdir, 'pages.tsv')
        ps1 = os.path.join(tmpdir, 'probe_word_pages.ps1')
        script = r'''
param(
  [Parameter(Mandatory=$true)][string]$DocxPath,
  [Parameter(Mandatory=$true)][string]$BookmarkList,
  [Parameter(Mandatory=$true)][string]$OutPath
)
$ErrorActionPreference = "Stop"
$word = New-Object -ComObject Word.Application
$word.Visible = $false
$word.DisplayAlerts = 0
$lines = New-Object System.Collections.Generic.List[string]
try {
  $doc = $word.Documents.Open($DocxPath, $false, $true)
  $doc.Repaginate()
  $names = Get-Content -LiteralPath $BookmarkList -Encoding UTF8
  foreach ($bm in $names) {
    try {
      $range = $doc.Bookmarks.Item($bm).Range
      $adj = $range.Information(1)
      $raw = $range.Information(3)
      $lines.Add("$bm`t$adj`t$raw")
    }
    catch {
      $msg = $_.Exception.Message -replace "(`r`n|`n|`r)", " "
      $lines.Add("$bm`tERR`t$msg")
    }
  }
  $doc.Close($false)
}
finally {
  $word.Quit()
}
[System.IO.File]::WriteAllLines($OutPath, $lines, [System.Text.UTF8Encoding]::new($false))
'''
        try:
            shutil.copy2(docx_path, probe_docx)
            remove_update_fields_from_docx(probe_docx)
            with open(list_path, 'w', encoding='utf-8') as f:
                f.write('\n'.join(unique_names))
            with open(ps1, 'w', encoding='utf-8') as f:
                f.write(script)
            subprocess.run(
                [
                    'powershell',
                    '-NoProfile',
                    '-ExecutionPolicy',
                    'Bypass',
                    '-File',
                    ps1,
                    '-DocxPath',
                    probe_docx,
                    '-BookmarkList',
                    list_path,
                    '-OutPath',
                    out_path,
                ],
                check=True,
                timeout=240,
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL,
            )
            raw = open(out_path, encoding='utf-8').read()
        finally:
            shutil.rmtree(tmpdir, ignore_errors=True)

        pages = {}
        errors = []
        for line in raw.splitlines():
            parts = line.split('\t')
            if len(parts) < 3:
                continue
            bm_name, adj, raw_page = parts[0], parts[1], parts[2]
            if adj == 'ERR':
                errors.append((bm_name, raw_page))
                continue
            pages[bm_name] = {'adjusted': adj, 'raw': raw_page}
        if errors:
            print(f'  [warn] Word bookmark 페이지 미검출 {len(errors)}개: {errors[:3]}')
        return pages

    probe_docx = os.path.join(
        os.path.dirname(os.path.abspath(docx_path)),
        f'toc_probe_{os.getpid()}.docx',
    )
    try:
        shutil.copy2(docx_path, probe_docx)
        remove_update_fields_from_docx(probe_docx)
        bm_list = ', '.join(_applescript_quote(name) for name in unique_names)
        script = f'''
set docPath to {_applescript_quote(probe_docx)}
tell application "Microsoft Word"
    open file name docPath read only true add to recent files false
    set theDoc to active document
    repaginate theDoc
    set bmNames to {{{bm_list}}}
    set out to ""
    repeat with bmName in bmNames
        try
            set br to text object of bookmark (bmName as text) of theDoc
            set adj to get range information br information type active end adjusted page number
            set rawp to get range information br information type active end page number
            set out to out & (bmName as text) & tab & adj & tab & rawp & linefeed
        on error errMsg
            set out to out & (bmName as text) & tab & "ERR" & tab & errMsg & linefeed
        end try
    end repeat
    close theDoc saving no
    return out
end tell
'''
        raw = subprocess.check_output(
            ['osascript'],
            input=script,
            text=True,
            stderr=subprocess.STDOUT,
            timeout=180,
        )
    finally:
        if os.path.exists(probe_docx):
            os.remove(probe_docx)

    pages = {}
    errors = []
    for line in raw.splitlines():
        parts = line.split('\t')
        if len(parts) < 3:
            continue
        bm_name, adj, raw_page = parts[0], parts[1], parts[2]
        if adj == 'ERR':
            errors.append((bm_name, raw_page))
            continue
        pages[bm_name] = {'adjusted': adj, 'raw': raw_page}
    if errors:
        print(f'  [warn] Word bookmark 페이지 미검출 {len(errors)}개: {errors[:3]}')
    return pages


def _extract_static_toc_numbers_from_docx(docx_path, labels):
    from zipfile import ZipFile
    from lxml import etree

    wanted = {_toc_static_key(label) for label in labels}
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    with ZipFile(docx_path, 'r') as zin:
        document_xml = zin.read('word/document.xml')
    root = etree.fromstring(document_xml)

    numbers = {}
    for p in root.xpath('//w:p', namespaces=ns):
        label, num = _toc_label_from_static_paragraph(p, ns)
        if label and _toc_static_key(label) in wanted:
            numbers[_toc_static_key(label)] = num
    return numbers


def freeze_toc_numbers_via_word(docx_path):
    """Microsoft Word가 DOCX 화면에서 계산한 페이지 번호로 목차를 고정한다."""
    pairs = _extract_toc_pageref_targets_from_docx(docx_path)
    if not pairs:
        print('  [warn] Word 목차 고정 실패: PAGEREF 목차 항목 없음')
        return False

    for attempt in range(3):
        word_pages = _word_bookmark_page_numbers(
            docx_path,
            [bm_name for _label, bm_name in pairs],
        )
        toc_numbers = {}
        raw_pages = {}
        for label, bm_name in pairs:
            page_info = word_pages.get(bm_name)
            if not page_info:
                continue
            toc_numbers[_toc_static_key(label)] = page_info['adjusted']
            raw_pages[_toc_static_key(label)] = page_info['raw']
        if len(toc_numbers) != len(pairs):
            print(f'  [warn] Word 목차 고정 실패: {len(toc_numbers)}/{len(pairs)}개만 산출')
            return False

        if not apply_toc_numbers_to_docx(docx_path, toc_numbers):
            return False
        sanitize_final_docx_fields(docx_path)

        visible_numbers = _extract_static_toc_numbers_from_docx(
            docx_path,
            [label for label, _bm_name in pairs],
        )
        mismatches = []
        for label, _bm_name in pairs:
            key = _toc_static_key(label)
            if visible_numbers.get(key) != toc_numbers.get(key):
                mismatches.append((label, visible_numbers.get(key), toc_numbers.get(key)))
        if not mismatches:
            print(f'  ✓ Word DOCX 기준 목차 페이지 번호 검증 통과: {len(pairs)}개 항목')
            chapter_keys = [
                _toc_static_key(label)
                for label, _bm in pairs
                if re.match(r'^제[0-9]+장\s+', label)
            ]
            chapter_summary = [
                f'{label}: {toc_numbers[_toc_static_key(label)]}'
                for label, _bm in pairs
                if _toc_static_key(label) in chapter_keys
            ]
            if chapter_summary:
                print('    장별 Word 표시 페이지: ' + ', '.join(chapter_summary))
            return True

        if attempt == 2:
            print(f'  [warn] Word 목차 검증 불일치 {len(mismatches)}개: {mismatches[:5]}')
            return False
        print('  🔁 Word 기준 목차 번호 재확인 후 갱신')

    return False


def _toc_label_before_pageref(paragraph, ns):
    parts = []
    for run in paragraph.findall('w:r', ns):
        fld = run.find('w:fldChar', ns)
        if fld is not None and fld.get(qn('w:fldCharType')) == 'begin':
            break
        parts.extend(run.xpath('.//w:t/text()', namespaces=ns))
    return ''.join(parts).strip()


def _paragraph_text(paragraph, ns):
    parts = []
    for el in paragraph.iter():
        if el.tag == qn('w:t'):
            parts.append(el.text or '')
        elif el.tag == qn('w:tab'):
            parts.append('\t')
        elif el.tag == qn('w:br'):
            parts.append('\n')
    return ''.join(parts)


def _replace_static_toc_number(paragraph, ns, old_num, new_num):
    text_nodes = paragraph.xpath('.//w:t', namespaces=ns)
    for node in reversed(text_nodes):
        text = node.text or ''
        replaced = re.sub(
            rf'({re.escape(str(old_num))})\s*$',
            str(new_num),
            text,
            count=1,
        )
        if replaced != text:
            node.text = replaced
            return True
    return False


def _toc_label_from_static_paragraph(paragraph, ns):
    text = _paragraph_text(paragraph, ns).strip()
    if not text:
        return None, None

    if '\t' in text:
        label, num = text.rsplit('\t', 1)
        label = re.sub(r'\.{2,}', ' ', label).strip()
        num = num.strip()
        if not re.match(r'^(?:[0-9]+|[ivxlcdmIVXLCDM]+)$', num):
            return None, None
    else:
        cleaned = re.sub(r'\.{2,}', ' ', text)
        m = re.match(
            r'^(?P<label>.+?)\s{2,}(?P<num>[0-9]+|[ivxlcdmIVXLCDM]+)\s*$',
            cleaned,
        )
        if not m:
            return None, None
        label = m.group('label').strip()
        num = m.group('num')

    if not (
        label.startswith('제')
        or label.startswith('<표')
        or label.startswith('<그림')
        or label in ('참고문헌', 'Abstract')
    ):
        return None, None
    return label, num


def apply_toc_numbers_to_docx(docx_path, toc_numbers):
    """PAGEREF 필드 상태와 이미 고정된 텍스트 상태를 모두 지원해 목차 숫자를 갱신한다."""
    import copy
    import tempfile
    from zipfile import ZipFile, ZIP_DEFLATED
    from lxml import etree

    if not toc_numbers:
        return False

    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    with ZipFile(docx_path, 'r') as zin:
        document_xml = zin.read('word/document.xml')

    root = etree.fromstring(document_xml)
    field_paragraphs = root.xpath(
        '//w:p[.//w:instrText[contains(., "PAGEREF")]]',
        namespaces=ns,
    )

    changed = 0
    missing = []
    handled = set()
    for p in field_paragraphs:
        label = _toc_label_before_pageref(p, ns)
        label_key = _toc_static_key(label)
        num = toc_numbers.get(label_key)
        if num is None:
            missing.append(label)
            continue

        runs = p.findall('w:r', ns)
        begin_idx = end_idx = None
        for idx, run in enumerate(runs):
            fld = run.find('w:fldChar', ns)
            if fld is not None and fld.get(qn('w:fldCharType')) == 'begin':
                begin_idx = idx
                break
        if begin_idx is None:
            missing.append(label)
            continue
        for idx in range(begin_idx, len(runs)):
            fld = runs[idx].find('w:fldChar', ns)
            if fld is not None and fld.get(qn('w:fldCharType')) == 'end':
                end_idx = idx
                break
        if end_idx is None:
            missing.append(label)
            continue

        r_new = etree.Element(qn('w:r'))
        rpr = None
        for candidate in runs[begin_idx:end_idx + 1] + list(reversed(runs[:begin_idx])):
            cand_rpr = candidate.find('w:rPr', ns)
            if cand_rpr is not None:
                rpr = copy.deepcopy(cand_rpr)
                break
        if rpr is not None:
            r_new.append(rpr)
        t = etree.SubElement(r_new, qn('w:t'))
        t.text = str(num)

        insert_at = p.index(runs[begin_idx])
        for old in runs[begin_idx:end_idx + 1]:
            p.remove(old)
        p.insert(insert_at, r_new)
        changed += 1
        handled.add(label_key)

    for p in root.xpath('//w:p', namespaces=ns):
        label, old_num = _toc_label_from_static_paragraph(p, ns)
        if not label:
            continue
        label_key = _toc_static_key(label)
        if label_key in handled:
            continue
        num = toc_numbers.get(label_key)
        if num is None:
            continue
        if str(old_num) != str(num) and _replace_static_toc_number(p, ns, old_num, num):
            changed += 1
        handled.add(label_key)

    if missing:
        print(f'  [warn] 목차 번호 고정 누락 {len(missing)}개: {missing[:3]}')
    missing_static = sorted(set(toc_numbers.keys()) - handled)
    if missing_static:
        print(f'  [warn] 목차 정적 항목 미검출 {len(missing_static)}개: {missing_static[:3]}')

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    tmp.close()
    try:
        with ZipFile(docx_path, 'r') as zin, ZipFile(tmp.name, 'w', ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                if item.filename == 'word/document.xml':
                    data = etree.tostring(
                        root,
                        xml_declaration=True,
                        encoding='UTF-8',
                        standalone=True,
                    )
                else:
                    data = zin.read(item.filename)
                zout.writestr(item, data)
        os.replace(tmp.name, docx_path)
    finally:
        if os.path.exists(tmp.name):
            os.remove(tmp.name)

    print(f'  ✓ 목차 페이지 번호 텍스트 고정/갱신: {changed}개')
    return not missing and not missing_static


def freeze_toc_numbers_from_pdf(docx_path, pdf_path):
    """목차/표목차/그림목차 PAGEREF를 최종 페이지 번호 텍스트로 고정한다.

    Word가 PDF 출력 시에는 PAGEREF를 계산하지만, DOCX 화면에서는 필드 cache가 1로
    보이는 경우가 있어 최종본은 필드 의존성을 제거한다.
    """
    toc_numbers = _extract_toc_numbers_from_pdf(pdf_path)
    if not toc_numbers:
        print('  [warn] 목차 번호 고정 생략: PDF에서 목차 번호를 추출하지 못함')
        return False
    return apply_toc_numbers_to_docx(docx_path, toc_numbers)


def freeze_toc_numbers_via_temp_pdf(docx_path):
    """최종 PDF 파일을 남기지 않고 DOCX 목차 PAGEREF 번호만 고정한다."""
    import shutil
    import tempfile

    with tempfile.TemporaryDirectory(prefix='thesis_toc_pdf_') as tmp:
        tmp_docx = os.path.join(tmp, os.path.basename(docx_path))
        shutil.copy2(docx_path, tmp_docx)
        freeze_fields_via_libreoffice(tmp_docx)
        if not export_pdf_via_libreoffice(tmp_docx):
            print('  [warn] 목차 번호 고정 실패: 임시 PDF 생성 실패')
            return False
        tmp_pdf = tmp_docx.replace('.docx', '.pdf')
        if not os.path.exists(tmp_pdf):
            print('  [warn] 목차 번호 고정 실패: 임시 PDF 파일 없음')
            return False
        return freeze_toc_numbers_from_pdf(docx_path, tmp_pdf)


def export_review_pdf(docx_path, final_pdf):
    use_word_pdf = os.environ.get('THESIS_USE_WORD_PDF') == '1'
    if use_word_pdf and export_pdf_via_word(docx_path, final_pdf):
        return os.path.exists(final_pdf)
    if use_word_pdf:
        print('  Word 변환 실패 → LibreOffice fallback')
    export_pdf_via_libreoffice(docx_path)
    return os.path.exists(final_pdf)


def main():
    print("📄 논문 DOCX v2 생성 (한양대 공식 서식)")

    doc = create_doc()
    enable_update_fields(doc)
    
    # 1. 표지 / 2. 제출서 / 3. 인준서 — 페이지번호 미표시 (한양대 표준)
    add_cover(doc)
    add_submission(doc)
    add_approval(doc)

    # 4. 목차/표목차/그림목차/국문요지 — 새 섹션, 로마자 i부터
    start_front_matter_section(doc)

    # 5. 본문 (제1장 만나면 다시 새 섹션 시작 → 아라비아 1부터)
    md_path = os.path.join(PAPER_DIR, '석사학위논문_박현근.md')
    bm_map = prebuild_bookmark_map(md_path)
    convert_md(doc, md_path, bm_map=bm_map)
    add_research_ethics_pledges(doc)
    
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    out = os.path.join(PAPER_DIR, f'석사학위논문_박현근_{timestamp}.docx')
    doc.save(out)

    if os.environ.get('THESIS_SKIP_TOC_FREEZE') == '1':
        print('🔧 목차 페이지 번호 고정 생략 (THESIS_SKIP_TOC_FREEZE=1)')
    elif os.environ.get('THESIS_TOC_SOURCE') == 'pdf':
        print('🔧 목차 페이지 번호 고정 (PDF 렌더 기준)')
        freeze_toc_numbers_via_temp_pdf(out)
    else:
        print('🔧 목차 페이지 번호 고정 (Microsoft Word DOCX 화면 기준)')
        if not freeze_toc_numbers_via_word(out):
            raise RuntimeError('Word DOCX 기준 목차 페이지 번호 고정 실패')

    export_pdf = os.environ.get('THESIS_EXPORT_PDF') == '1'
    if export_pdf:
        # PDF 검수는 명시 opt-in이다. 제출 기준은 DOCX이며, PDF 변환은 렌더러별 폰트 차이가 있다.
        print('🔧 PDF 생성 (검수용, opt-in)')
        final_pdf = out.replace('.docx', '.pdf')
        if not export_review_pdf(out, final_pdf):
            raise RuntimeError('PDF 생성 실패')
        if os.environ.get('THESIS_VALIDATE_PDF_TOC') == '1':
            ok, _report = validate_toc_numbers_in_pdf(final_pdf)
            if not ok:
                raise RuntimeError('최종 PDF 목차 페이지 번호 검증 실패')
        else:
            print('  ✓ PDF 생성 완료; 목차 검증 기준은 Word DOCX 화면')
    else:
        print('🔧 PDF 생성 생략 (DOCX 제출용; 필요 시 THESIS_EXPORT_PDF=1)')

    print(f"✅ 완료: {out}")
    print(f"   크기: {os.path.getsize(out)//1024}KB")
    print(f"   문단: {len(doc.paragraphs)}, 표: {len(doc.tables)}")
    print(f"   여백: 상하 4cm, 좌우 3.5cm (공식 워드 기준)")
    print(f"   폰트: {FONT_KR} 11pt (전체 통일), 표/그림 10pt")


if __name__ == '__main__':
    main()
