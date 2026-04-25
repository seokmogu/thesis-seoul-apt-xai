#!/usr/bin/env python3
"""
논문 DOCX 변환 v2 — 한양대학교 부동산융합대학원 석사논문 공식 서식
출처: https://gupd.hanyang.ac.kr/front/ko/bachelor/thesis
참고: 조민지(2023) 한양대 석사논문

편집용지: A4
워드 여백: 위 4cm, 아래 4cm, 좌 3.5cm, 우 3.5cm, 머리말 1.5cm, 꼬리말 1.5cm
본문: 10~11pt 바탕체, 줄간격 160~200%
큰제목(장): 16pt 진하게 / 중간제목(절): 13pt 진하게
"""
import os, re, io
import latex2mathml.converter as _l2m
import mathml2omml as _m2o
from docx import Document

OMML_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'

from docx.shared import Pt, Cm, Mm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

BASE_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), '..')
PAPER_DIR = os.path.join(BASE_DIR, 'paper')

FONT_KR = '바탕'


def set_font(run, size=11, bold=False, font=FONT_KR):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{font}"/>')
        rPr.append(rFonts)
    else:
        rFonts.set(qn('w:eastAsia'), font)


def set_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_page_number(section, start_num=None, fmt='decimal'):
    """페이지 번호 추가 (하단 중앙, "- N -" 형식)"""
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    
    r1 = p.add_run('- ')
    set_font(r1, 10)
    
    # 페이지 번호 필드
    fld_xml = (
        f'<w:fldSimple {nsdecls("w")} w:instr=" PAGE \\* {fmt} ">'
        f'<w:r><w:rPr><w:sz w:val="20"/></w:rPr><w:t>1</w:t></w:r>'
        f'</w:fldSimple>'
    )
    p._element.append(parse_xml(fld_xml))
    
    r2 = p.add_run(' -')
    set_font(r2, 10)


def create_doc():
    doc = Document()
    
    # 기본 섹션 설정
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Cm(4)
    section.bottom_margin = Cm(4)
    section.left_margin = Cm(3.5)
    section.right_margin = Cm(3.5)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.5)
    
    # 기본 스타일
    style = doc.styles['Normal']
    style.font.name = FONT_KR
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    style.paragraph_format.line_spacing = 1.6
    style.paragraph_format.first_line_indent = Cm(0.8)
    rPr = style.element.get_or_add_rPr()
    rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{FONT_KR}"/>')
    rPr.append(rFonts)
    
    return doc


def latex_to_omml(latex, block=False):
    """LaTeX -> MathML -> OMML XML element (with m: namespace). block=True wraps in oMathPara."""
    mml = _l2m.convert(latex)
    omml_str = _m2o.convert(mml)  # starts with <m:oMath>...</m:oMath>
    # Bug workaround: mathml2omml 0.0.2 emits malformed closing tag for groupChrPr.
    omml_str = re.sub(
        r'(<m:groupChrPr>[^<]*<m:chr[^/]*/>[^<]*<m:pos[^/]*/>)\s*</m:groupChr>',
        r'\1</m:groupChrPr>',
        omml_str,
    )
    # Inject namespace on root element so parse_xml can resolve m: prefix.
    omml_ns = omml_str.replace('<m:oMath>', f'<m:oMath xmlns:m="{OMML_NS}">', 1)
    if block:
        omml_ns = (
            f'<m:oMathPara xmlns:m="{OMML_NS}">'
            f'{omml_ns}'
            f'</m:oMathPara>'
        )
    return parse_xml(omml_ns)


def add_block_math(doc, latex):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    try:
        elem = latex_to_omml(latex, block=True)
        p._p.append(elem)
    except Exception as e:
        r = p.add_run(latex)
        r.font.size = Pt(10)
        r.font.name = 'Cambria Math'
        r.italic = True
        print(f'  [math fallback-block] {latex[:60]} -> {e}')
    return p


def add_inline_math_run(paragraph, latex):
    try:
        elem = latex_to_omml(latex, block=False)
        paragraph._p.append(elem)
        return None
    except Exception as e:
        r = paragraph.add_run(latex)
        r.font.name = 'Cambria Math'
        r.italic = True
        print(f'  [math fallback-inline] {latex[:60]} -> {e}')
        return r


def centered_text(doc, text, size=14, bold=False, space_before=0, space_after=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    if space_before:
        p.paragraph_format.space_before = Pt(space_before)
    if space_after:
        p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    set_font(r, size, bold)
    return p


def add_cover(doc):
    """양식1: 표지"""
    for _ in range(4):
        doc.add_paragraph('')
    centered_text(doc, '석 사 학 위 논 문', 16, True)
    for _ in range(2):
        doc.add_paragraph('')
    centered_text(doc, 'XGBoost와 SHAP을 활용한', 22, True)
    centered_text(doc, '서울시 아파트 매매가격 결정요인 분석', 22, True)
    doc.add_paragraph('')
    centered_text(doc, 'Analysis of Determinants of Apartment Sale Prices', 14, True)
    centered_text(doc, 'in Seoul Using XGBoost and SHAP', 14, True)
    for _ in range(5):
        doc.add_paragraph('')
    centered_text(doc, '박  현  근', 16, True)
    for _ in range(2):
        doc.add_paragraph('')
    centered_text(doc, '한 양 대 학 교  부 동 산 융 합 대 학 원', 14, True)
    doc.add_paragraph('')
    centered_text(doc, '2026 년  2 월', 14, True)
    doc.add_page_break()


def add_submission(doc):
    """양식2: 제출서"""
    for _ in range(4):
        doc.add_paragraph('')
    centered_text(doc, '석 사 학 위 논 문', 16, True)
    for _ in range(2):
        doc.add_paragraph('')
    centered_text(doc, 'XGBoost와 SHAP을 활용한', 22, True)
    centered_text(doc, '서울시 아파트 매매가격 결정요인 분석', 22, True)
    doc.add_paragraph('')
    centered_text(doc, 'Analysis of Determinants of Apartment Sale Prices', 14, True)
    centered_text(doc, 'in Seoul Using XGBoost and SHAP', 14, True)
    for _ in range(2):
        doc.add_paragraph('')
    centered_text(doc, '지도교수  ____________', 14, True)
    for _ in range(2):
        doc.add_paragraph('')
    centered_text(doc, '이 논문을 공학 석사학위논문으로 제출합니다.', 12)
    for _ in range(2):
        doc.add_paragraph('')
    centered_text(doc, '2026 년  2 월', 14, True)
    doc.add_paragraph('')
    centered_text(doc, '한 양 대 학 교  부 동 산 융 합 대 학 원', 14, True)
    doc.add_paragraph('')
    centered_text(doc, '빅 데 이 터  전 공', 13)
    doc.add_paragraph('')
    centered_text(doc, '박  현  근', 16, True)
    doc.add_page_break()


def add_approval(doc):
    """양식3: 인준서"""
    for _ in range(3):
        doc.add_paragraph('')
    centered_text(doc, '이 논문을 박현근의', 14)
    centered_text(doc, '석사학위 논문으로 인준함.', 14)
    for _ in range(2):
        doc.add_paragraph('')
    centered_text(doc, '2026 년  2 월', 14, True)
    for _ in range(4):
        doc.add_paragraph('')
    
    for role in ['심 사 위 원 장', '심  사  위  원', '심  사  위  원']:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_after = Pt(24)
        r = p.add_run(f'{role} :  ________________  (인)')
        set_font(r, 14)
    
    for _ in range(3):
        doc.add_paragraph('')
    centered_text(doc, '한 양 대 학 교  부 동 산 융 합 대 학 원', 14, True)
    doc.add_page_break()


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(24)
        p.paragraph_format.space_after = Pt(18)
        r = p.add_run(text)
        set_font(r, 16, True)
    elif level == 2:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(12)
        r = p.add_run(text)
        set_font(r, 13, True)
    elif level == 3:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(8)
        r = p.add_run(text)
        set_font(r, 11, True)
    elif level == 4:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(text)
        set_font(r, 12, True)
    return p


def add_body(doc, text):
    if not text.strip():
        return
    p = doc.add_paragraph()
    # split by bold first
    bold_parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for bp in bold_parts:
        is_bold = bp.startswith('**') and bp.endswith('**')
        segment = bp[2:-2] if is_bold else bp
        # split by inline math $...$
        math_parts = re.split(r'(\$[^$\n]+\$)', segment)
        for mp in math_parts:
            if mp.startswith('$') and mp.endswith('$') and len(mp) > 2:
                add_inline_math_run(p, mp[1:-1])
            elif mp:
                r = p.add_run(mp)
                set_font(r, 11, is_bold)
    return p


def is_sep(line):
    return bool(re.match(r'^\s*\|[\s:-]+\|\s*$', line))


def parse_row(line):
    return [c.strip() for c in line.strip().strip('|').split('|')]


def add_table(doc, rows, caption=''):
    if caption:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_before = Pt(8)
        r = p.add_run(caption)
        set_font(r, 10, True, '휴먼명조')  # 공식: 표/그림 내용 10pt 휴먼명조
    
    if not rows:
        return
    
    ncols = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=ncols)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for i, row_data in enumerate(rows):
        for j, txt in enumerate(row_data):
            if j >= ncols:
                continue
            cell = tbl.rows[i].cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            r = p.add_run(txt)
            set_font(r, 9 if i > 0 else 9, bold=(i == 0), font='휴먼명조')
            if i == 0:
                set_shading(cell, 'D9E2F3')
    
    doc.add_paragraph('')


def add_image(doc, path, caption=''):
    if not os.path.exists(path):
        add_body(doc, f'[그림 파일 없음: {os.path.basename(path)}]')
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run()
    r.add_picture(path, width=Cm(13))
    
    if caption:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        r = p.add_run(caption)
        set_font(r, 10, font='휴먼명조')


def convert_md(doc, md_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    start = 0
    for i, line in enumerate(lines):
        if line.strip().startswith('# 국문초록') or line.strip().startswith('# 국문 초록'):
            start = i
            break
    
    i = start
    caption = ''
    
    while i < len(lines):
        line = lines[i].rstrip()
        
        # 블록 수식 (한 줄 $$...$$)
        m_one = re.match(r'^\s*\$\$(.+)\$\$\s*$', line)
        if m_one:
            add_block_math(doc, m_one.group(1).strip())
            i += 1
            continue
        # 블록 수식 (다중 라인 $$ ... $$)
        if line.strip() == '$$':
            parts = []
            i += 1
            while i < len(lines) and lines[i].strip() != '$$':
                parts.append(lines[i].strip())
                i += 1
            add_block_math(doc, ' '.join(parts))
            i += 1
            continue
        
        # 표 캡션
        m = re.match(r'^\*\*(<표[^>]*>[^*]*)\*\*$', line.strip())
        if m:
            caption = m.group(1)
            i += 1
            continue
        
        # 표
        if '|' in line and line.strip().startswith('|') and not is_sep(line):
            rows = []
            while i < len(lines):
                l = lines[i].rstrip()
                if '|' in l and l.strip().startswith('|'):
                    if not is_sep(l):
                        rows.append(parse_row(l))
                    i += 1
                else:
                    break
            add_table(doc, rows, caption)
            caption = ''
            continue
        
        if is_sep(line):
            i += 1
            continue
        
        # 그림
        m = re.match(r'!\[(.+?)\]\((.+?)\)', line)
        if m:
            cap = m.group(1)
            rel = m.group(2)
            absp = os.path.normpath(os.path.join(PAPER_DIR, rel))
            add_image(doc, absp, cap)
            i += 1
            continue
        
        # 제목
        if line.startswith('# '):
            txt = line[2:].strip()
            # 장 제목은 새 페이지
            if txt.startswith('제') and '장' in txt:
                doc.add_page_break()
            add_heading(doc, txt, 1)
            i += 1
            continue
        if line.startswith('## '):
            add_heading(doc, line[3:].strip(), 2)
            i += 1
            continue
        if line.startswith('#### '):
            add_heading(doc, line[5:].strip(), 4)
            i += 1
            continue
        if line.startswith('### '):
            add_heading(doc, line[4:].strip(), 3)
            i += 1
            continue
        
        if line.strip() == '---':
            doc.add_page_break()
            i += 1
            continue
        
        if line.strip().startswith('- <표') or line.strip().startswith('- <그림'):
            i += 1
            continue
        
        if not line.strip():
            i += 1
            continue
        
        add_body(doc, line)
        i += 1


def main():
    print("📄 논문 DOCX v2 생성 (한양대 공식 서식)")
    
    doc = create_doc()
    
    # 1. 표지
    add_cover(doc)
    
    # 2. 제출서
    add_submission(doc)
    
    # 3. 인준서
    add_approval(doc)
    
    # 4. 본문 (목차~참고문헌~Abstract 포함)
    convert_md(doc, os.path.join(PAPER_DIR, '논문_초안.md'))
    
    # 페이지 번호 추가 (첫 번째 섹션)
    add_page_number(doc.sections[0])
    
    out = os.path.join(PAPER_DIR, '논문_초안.docx')
    doc.save(out)
    
    print(f"✅ 완료: {out}")
    print(f"   크기: {os.path.getsize(out)//1024}KB")
    print(f"   문단: {len(doc.paragraphs)}, 표: {len(doc.tables)}")
    print(f"   여백: 상하 4cm, 좌우 3.5cm (공식 워드 기준)")
    print(f"   폰트: 바탕 11pt, 표/그림 휴먼명조 10pt")


if __name__ == '__main__':
    main()
