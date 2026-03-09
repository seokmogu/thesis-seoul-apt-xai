#!/usr/bin/env python3
"""
논문 DOCX 변환 — 한양대학교 부동산융합대학원 석사논문 형식
- 편집용지: A4, 여백 위아래 38mm, 좌우 35mm
- 본문: 바탕체 11pt, 줄간격 160%
- 큰제목: 16pt 진하게, 중간제목: 13pt 진하게
"""
import os, re, sys
from docx import Document
from docx.shared import Pt, Cm, Mm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

BASE_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), '..')
PAPER_DIR = os.path.join(BASE_DIR, 'paper')
FIGURES_DIR = os.path.join(BASE_DIR, 'figures')
RESULTS_DIR = os.path.join(BASE_DIR, 'results')

FONT_NAME = '바탕'
FONT_NAME_EA = '바탕'


def set_run_font(run, size=11, bold=False, name=FONT_NAME):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{FONT_NAME_EA}"/>')
        rPr.append(rFonts)
    else:
        rFonts.set(qn('w:eastAsia'), FONT_NAME_EA)


def set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def create_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(38)
    section.bottom_margin = Mm(38)
    section.left_margin = Mm(35)
    section.right_margin = Mm(35)
    section.header_distance = Mm(15)
    section.footer_distance = Mm(15)

    style = doc.styles['Normal']
    style.font.name = FONT_NAME
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    style.paragraph_format.line_spacing = 1.6
    style.paragraph_format.first_line_indent = Cm(0.8)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{FONT_NAME_EA}"/>')
        rPr.append(rFonts)
    else:
        rFonts.set(qn('w:eastAsia'), FONT_NAME_EA)
    return doc


def add_cover(doc):
    for _ in range(6):
        doc.add_paragraph('')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('석사학위논문')
    set_run_font(r, 16)
    doc.add_paragraph('')
    doc.add_paragraph('')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('XGBoost와 SHAP을 활용한\n서울시 아파트 매매가격 결정요인 분석')
    set_run_font(r, 22, bold=True)
    for _ in range(6):
        doc.add_paragraph('')
    for text in ['2026년  2월', '한양대학교 부동산융합대학원']:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        set_run_font(r, 14)
        doc.add_paragraph('')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('[성 명]')
    set_run_font(r, 16, bold=True)
    doc.add_page_break()


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    p.paragraph_format.space_after = Pt(12 if level == 1 else 8)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    sizes = {1: 16, 2: 13, 3: 11}
    r = p.add_run(text)
    set_run_font(r, sizes.get(level, 11), bold=True)
    return p


def add_body(doc, text):
    if not text.strip():
        return
    p = doc.add_paragraph()
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            r = p.add_run(part[2:-2])
            set_run_font(r, 11, bold=True)
        else:
            r = p.add_run(part)
            set_run_font(r, 11)
    return p


def add_table(doc, rows_data, caption=''):
    if caption:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_before = Pt(6)
        r = p.add_run(caption)
        set_run_font(r, 10, bold=True)

    if not rows_data:
        return

    n_cols = max(len(r) for r in rows_data)
    tbl = doc.add_table(rows=len(rows_data), cols=n_cols)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row in enumerate(rows_data):
        for j, cell_text in enumerate(row):
            if j >= n_cols:
                continue
            cell = tbl.rows[i].cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            r = p.add_run(cell_text)
            set_run_font(r, 9, bold=(i == 0))
            if i == 0:
                set_cell_shading(cell, 'D9E2F3')

    p = doc.add_paragraph('')
    p.paragraph_format.space_before = Pt(4)


def add_image(doc, img_path, caption=''):
    if not os.path.exists(img_path):
        add_body(doc, f'[그림 파일 없음: {os.path.basename(img_path)}]')
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run()
    r.add_picture(img_path, width=Cm(14))
    if caption:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        r = p.add_run(caption)
        set_run_font(r, 10)


def is_separator(line):
    return bool(re.match(r'^\s*\|[\s:-]+\|\s*$', line))


def parse_table_row(line):
    return [c.strip() for c in line.strip().strip('|').split('|')]


def convert(doc, md_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    # 국문초록부터 시작
    start_idx = 0
    for i, line in enumerate(lines):
        if line.strip().startswith('# 국문초록') or line.strip().startswith('# 국문 초록'):
            start_idx = i
            break

    i = start_idx
    pending_caption = ''

    while i < len(lines):
        line = lines[i].rstrip()

        # 수식 블록
        if line.strip() == '$$':
            formula_lines = []
            i += 1
            while i < len(lines) and lines[i].strip() != '$$':
                formula_lines.append(lines[i].strip())
                i += 1
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Cm(0)
            r = p.add_run(' '.join(formula_lines))
            r.font.size = Pt(10)
            r.font.name = 'Cambria Math'
            r.italic = True
            i += 1
            continue

        # 표 캡션 감지
        cap_match = re.match(r'^\*\*(<표[^>]*>[^*]*)\*\*$', line.strip())
        if cap_match:
            pending_caption = cap_match.group(1)
            i += 1
            continue

        # 표 시작 감지
        if '|' in line and line.strip().startswith('|') and not is_separator(line):
            table_rows = []
            while i < len(lines):
                l = lines[i].rstrip()
                if '|' in l and l.strip().startswith('|'):
                    if not is_separator(l):
                        table_rows.append(parse_table_row(l))
                    i += 1
                else:
                    break
            add_table(doc, table_rows, pending_caption)
            pending_caption = ''
            continue

        # 구분선 스킵
        if is_separator(line):
            i += 1
            continue

        # 그림
        img_match = re.match(r'!\[(.+?)\]\((.+?)\)', line)
        if img_match:
            caption = img_match.group(1)
            rel_path = img_match.group(2)
            abs_path = os.path.normpath(os.path.join(PAPER_DIR, rel_path))
            add_image(doc, abs_path, caption)
            i += 1
            continue

        # 제목
        if line.startswith('# '):
            text = line[2:].strip()
            if text.startswith('제') and '장' in text:
                doc.add_page_break()
            add_heading(doc, text, 1)
            i += 1
            continue
        if line.startswith('## '):
            add_heading(doc, line[3:].strip(), 2)
            i += 1
            continue
        if line.startswith('### '):
            add_heading(doc, line[4:].strip(), 3)
            i += 1
            continue

        # --- 페이지 브레이크
        if line.strip() == '---':
            doc.add_page_break()
            i += 1
            continue

        # 목차 항목 스킵
        if line.strip().startswith('- <표') or line.strip().startswith('- <그림'):
            i += 1
            continue

        # 빈 줄 스킵
        if not line.strip():
            i += 1
            continue

        # 일반 본문
        add_body(doc, line)
        i += 1


def main():
    print("📄 논문 DOCX 생성 (한양대 석사논문 형식)")

    doc = create_document()
    add_cover(doc)
    convert(doc, os.path.join(PAPER_DIR, '논문_초안.md'))

    out = os.path.join(PAPER_DIR, '논문_초안.docx')
    doc.save(out)

    n_p = len(doc.paragraphs)
    n_t = len(doc.tables)
    sz = os.path.getsize(out) // 1024

    print(f"✅ 완료: {out}")
    print(f"   크기: {sz}KB, 문단: {n_p}, 표: {n_t}")
    print(f"   Word에서 열어 실제 페이지 수 확인하세요")


if __name__ == '__main__':
    main()
